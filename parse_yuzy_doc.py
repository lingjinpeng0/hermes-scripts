#!/usr/bin/env python3
"""
渔业资源学课件整理 — Parse detailed chapter markdowns into formatted .docx
Robust parser that handles bullet lists, numbered items, Chinese headings, tables.
"""
import os, re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")
SRC_DIR = os.path.expanduser(r"~\.hermes\project-data\渔业资源学")
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

CHAPTER_FILES = ['ch1_detailed.md','ch2_detailed.md','ch3_detailed.md','ch4_detailed.md','ch5_detailed.md']

# ── Font ──
def sf(run, cn='宋体', en='Times New Roman', sz=Pt(10.5), b=False, i=False):
    run.font.size = sz; run.font.name = en
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.bold = b; run.font.italic = i

def body(doc, text, indent=True):
    if not text: return
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r); p.paragraph_format.line_spacing=1.2
    p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def mk_h(doc, text, lv=1, num=''):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {lv}']
    display = f'{num} {text}' if num else text
    r = p.add_run(display)
    sizes = {1:Pt(16),2:Pt(14),3:Pt(12)}
    sf(r, cn='黑体', sz=sizes.get(lv,Pt(12)), b=False, i=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    sb = {1:Pt(18),2:Pt(12),3:Pt(8)}.get(lv,Pt(8))
    sa = {1:Pt(6),2:Pt(4),3:Pt(3)}.get(lv,Pt(3))
    p.paragraph_format.space_before=sb; p.paragraph_format.space_after=sa

# ── Table ──
def tbl(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); sf(r); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.first_line_indent=None
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tl = t._tbl; tp = tl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); tl.insert(0,tp)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct'); tp.append(tw)
    tb=OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'nil'); tb.append(b)
    tp.append(tb)
    for ci,hd in enumerate(headers):
        c=t.rows[0].cells[ci]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Pt(0)
        r=p.add_run(hd); sf(r,sz=Pt(10))
        _bdr(c,'12','6')
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if len(str(val))<30 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent=Pt(0)
            r=p.add_run(str(val)); sf(r,sz=Pt(9))
    for ci in range(len(headers)):
        c=t.rows[-1].cells[ci]; tcp=c._tc.get_or_add_tcPr()
        tcb=OxmlElement('w:tcBorders'); bt=OxmlElement('w:bottom')
        bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'12')
        bt.set(qn('w:space'),'0'); bt.set(qn('w:color'),'000000')
        tcb.append(bt); tcp.append(tcb)
    for row in tl.findall(qn('w:tr')):
        trp=row.find(qn('w:trPr'))
        if trp is None: trp=OxmlElement('w:trPr'); row.insert(0,trp)
        old=trp.find(qn('w:cantSplit'))
        if old is not None: trp.remove(old)
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trp.append(cs)

def _bdr(c,ts='12',bs='6'):
    tcp=c._tc.get_or_add_tcPr(); tcb=OxmlElement('w:tcBorders')
    for s,sz in [('top',ts),('bottom',bs)]:
        e=OxmlElement(f'w:{s}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); tcb.append(e)
    tcp.append(tcb)

# ── Heading detection ──
def hd_lvl(line):
    m=re.match(r'^(#{1,4})\s+(.+)$', line)
    return (len(m.group(1)), m.group(2).strip()) if m else (None,None)

def clean(text):
    text=re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text=re.sub(r'\*(.+?)\*', r'\1', text)
    text=re.sub(r'`([^`]+)`', r'\1', text)
    emoji=re.compile("[\U0001F600-\U0001F64F\U0001F300-\U0001F5FF\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF\U00002702-\U000027B0\U000024C2-\U0001F251\U0001F900-\U0001F9FF\U0001FA00-\U0001FA6F\U0001FA70-\U0001FAFF]+")
    return emoji.sub('', text)

# ── Process chapter ──
def proc_ch(doc, fpath, ch_num):
    with open(fpath,'r',encoding='utf-8') as f:
        src = f.read()
    # Clean up markdown bold markers globally
    src = re.sub(r'\*\*(.+?)\*\*', r'\1', src)
    # Convert horizontal rules
    src = re.sub(r'^---+$', '', src, flags=re.MULTILINE)
    lines = src.split('\n')

    in_qa = False; qa_n = 0
    tbl_buf = []; in_tbl = False
    h2_n = 0

    # Chapter title from first heading
    if lines:
        flvl, ftxt = hd_lvl(lines[0])
        if flvl == 1 and ftxt:
            txt = re.sub(r'^第[一二三四五六七八九十]+章\s*', '', ftxt)
            txt = re.sub(r'[─—–-]+.*$', '', txt).strip()
            mk_h(doc, txt, 1, f'{ch_num}.')

    for line in lines[1:]:
        s = line.rstrip()
        if not s:
            if in_tbl and tbl_buf:
                _flush_tbl(doc, tbl_buf); tbl_buf=[]; in_tbl=False
            continue

        # Image refs
        if s.startswith('![', 0, 3) or '![' in s[:3] and '](' in s:
            continue

        # Code fences
        if s.startswith('```') or s.startswith('~~~'):
            continue

        # Table lines
        if s.startswith('|') and s.count('|') >= 3:
            cells = [c.strip() for c in s.split('|')[1:-1]]
            cells = [c for c in cells if c]
            if len(cells) >= 2:
                # Check if this is a separator row
                if re.match(r'^[-:\s|]+$', s):
                    continue
                tbl_buf.append(cells); in_tbl = True
                continue
        else:
            if in_tbl and tbl_buf:
                _flush_tbl(doc, tbl_buf); tbl_buf = []; in_tbl = False

        # Headings
        lvl, txt = hd_lvl(s)
        if lvl is not None:
            txt = clean(txt)
            if '思考题' in txt and lvl <= 3:
                in_qa = True; qa_n = 0
                mk_h(doc, '思考题', 2, f'{ch_num}.99')
                continue
            if in_qa:
                qm = re.match(r'思考题[一二三四五六七八九十][：:]\s*(.*)', txt)
                if qm:
                    qa_n += 1
                    _qa_q(doc, qa_n, qm.group(1))
                    continue
                continue  # Skip other headings in QA section
            if lvl == 2:
                h2_n += 1
                sect = re.sub(r'^第[一二三四五六七八九十]+节\s*', '', txt)
                mk_h(doc, sect, 2, f'{ch_num}.{h2_n}')
            elif lvl in (3, 4):
                # Find sub-number
                subn = 1
                mk_h(doc, txt, 3, f'{ch_num}.{h2_n}.{subn}')
            continue

        # Body text
        cl = clean(s)
        if len(cl) < 6:
            continue

        if in_qa:
            if '答案' in cl:
                _qa_ans(doc, '答案:')
                # Check for inline numbered items
                rest = re.split(r'答案[\s]*[：:]\s*', cl, maxsplit=1)
                if len(rest) > 1:
                    _proc_qa_items(doc, rest[1])
                continue
            if '解析' in cl and (cl.startswith('解析') or cl.startswith('解析')):
                expl = re.sub(r'^解析[\s]*[：:]\s*', '', cl)
                _qa_exp(doc, expl)
                continue
            im = re.match(r'[\(（](\d+)[\)）]\s*(.+)', cl)
            if im:
                _qa_ap(doc, int(im.group(1)), im.group(2))
                continue
            _proc_qa_items(doc, cl)
            continue

        # Strip common prefixes
        cl = re.sub(r'^[-*]\s+', '', cl)
        if cl:
            body(doc, cl)

def _flush_tbl(doc, buf):
    if len(buf) < 2: return
    hd = buf[0]; n = len(hd)
    rows = []
    for r in buf[1:]:
        if len(r) >= 2:
            rr = r[:n] if len(r) >= n else r + ['']*(n-len(r))
            rows.append(rr)
    if rows and n >= 2:
        try: tbl(doc, hd, rows)
        except: pass

def _qa_q(doc, n, txt):
    p = doc.add_paragraph(); r = p.add_run(f'{n}. {txt}')
    sf(r, sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def _qa_ans(doc, txt):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21)
    r = p.add_run(txt); sf(r)

def _qa_ap(doc, n, txt):
    p = doc.add_paragraph()
    r = p.add_run(f'({n}) {txt}'); sf(r, sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr = p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210')
    pPr.append(ind)

def _qa_exp(doc, txt):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21)
    r = p.add_run(f'解析: {txt}'); sf(r, sz=Pt(10))

def _proc_qa_items(doc, text):
    items = re.split(r'(?=[\(（]\d+[\)）])', text)
    for it in items:
        it = it.strip()
        if not it: continue
        m = re.match(r'[\(（](\d+)[\)）]\s*(.*)', it)
        if m:
            _qa_ap(doc, int(m.group(1)), m.group(2))
        else:
            body(doc, it)

# ── MAIN ──
def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name='Times New Roman'; style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    style.font.size=Pt(10.5); style.font.bold=False; style.font.italic=False
    style.paragraph_format.line_spacing=1.2; style.paragraph_format.first_line_indent=Pt(21)

    for lv in [1,2,3]:
        s=doc.styles[f'Heading {lv}']
        s.font.name='Times New Roman'; s.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
        s.font.bold=False; s.font.italic=False; s.font.color.rgb=None

    # Cover
    for _ in range(6): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('渔业资源学'); sf(r,cn='黑体',sz=Pt(26)); p.paragraph_format.first_line_indent=None
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Fishery Resource Science'); sf(r,sz=Pt(14)); p.paragraph_format.first_line_indent=None
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('课件整理'); sf(r,cn='黑体',sz=Pt(18)); p.paragraph_format.first_line_indent=None
    doc.add_page_break()

    for cn, fn in enumerate(CHAPTER_FILES, 1):
        fp = os.path.join(SRC_DIR, fn)
        if not os.path.exists(fp): print(f'SKIP: {fn}'); continue
        print(f'Processing Ch{cn}...')
        proc_ch(doc, fp, cn)
        if cn < 5: doc.add_page_break()

    doc.save(OUTPUT)
    total=sum(len(p.text) for p in doc.paragraphs)
    print(f'文档已生成: {OUTPUT}')
    print(f'总字符数: {total}')
    print(f'表格数: {len(doc.tables)}')
    print(f'段落数: {len(doc.paragraphs)}')
    bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
    print(f'粗体运行: {bold_c}')

if __name__=='__main__':
    main()
