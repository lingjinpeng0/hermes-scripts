#!/usr/bin/env python3
"""Append Chapter 5 to existing docx."""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")
LQ='\u201c'; RQ='\u201d'

def sf(r, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False, italic=False):
    r.font.size=sz; r.font.name=en
    r.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.bold=bold; r.font.italic=italic

def body(doc, txt, indent=True):
    p=doc.add_paragraph(); r=p.add_run(txt); sf(r)
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def h1(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 1']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(16),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

def h2(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 2']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(14),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def h3(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 3']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(12),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def q(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'{n}. {txt}'); sf(r,sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def ans(doc):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run('答案:'); sf(r)

def ap(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'({n}) {txt}'); sf(r,sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210'); pPr.append(ind)

def exp(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run(f'解析: {txt}'); sf(r,sz=Pt(10))

doc = Document(DOCX)
doc.add_page_break()

h1(doc, '5. 渔业资源保护与管理')

h2(doc, '5.1 渔业资源相关权属')
body(doc, '渔业资源的权属制度是渔业管理的基础。根据《中华人民共和国宪法》和《渔业法》，渔业资源属于国家所有（全民所有）。管理权方面，实行统一领导、分级行政监督的管理体制，国务院渔业行政主管部门负责全国渔业监督管理工作。使用权方面，依法取得捕捞许可证的单位和个人享有在特定水域从事捕捞生产的权利，但必须遵守资源保护的管理规定。')

h2(doc, '5.2 新《渔业法》与渔业资源保护')
h3(doc, '5.2.1 《渔业法》的发展历程')
body(doc, '中国渔业法制建设经历了漫长的发展过程：1957年4月，《水产资源繁殖保护暂行条例（草案）》开始征求意见；1979年2月10日，国务院正式发布该条例（共8章20条）。1986年1月20日，第六届全国人大常委会第十四次会议通过《中华人民共和国渔业法》（共6章50条），标志着渔业管理进入法治化轨道。此后历经四次修正——2000年（增加18条、修改15条、删除3条）、2004年、2009年和2013年。')
body(doc, '2024至2026年的最新修订从6章50条扩展为7章90条。修订重点包括：(1)总则确立顶层发展导向与基本方针；(2)养殖业全流程规范，推动绿色生态养殖转型；(3)捕捞业双向管控强度，实现捕捞可持续；(4)渔业资源保护构建全方位生态保护体系；(5)增设独立的监督管理章节，强化综合监管能力；(6)法律责任大幅提高违法成本，强化震慑力度。')

h2(doc, '5.3 水生野生动物保护和管理')
body(doc, '水生野生动物保护是中国生物多样性保护的重要组成部分。《国家重点保护野生动物名录》涵盖从哺乳纲到珊瑚纲的多个分类阶元，保护对象包括白鱀豚、中华鲟、大鲵、海龟、砗磲等珍稀濒危物种。中国还签署了CITES（濒危野生动植物种国际贸易公约），加强水生野生动物的进出口管理。')
body(doc, '水生野生动物保护与管理的基本制度包括：(1)特许捕捉制度；(2)驯养繁殖许可证制度；(3)经营利用管理制度；(4)运输管理制度；(5)进出口管理制度；(6)保护补偿制度；(7)保护专项资金制度。主要保护措施包括建立自然保护区、实施禁渔期和禁渔区、划定重要栖息地、开展增殖放流、加强执法监管等。')

h2(doc, '5.4 渔业资源养护与管理技术与制度')
h3(doc, '5.4.1 渔业资源的概念与特性')
body(doc, '渔业资源的自然特性决定了管理的特殊性：(1)可再生性——在合理利用下能自我更新，但过度利用会导致资源衰退；(2)流动性——渔业资源分布广泛且不断移动，增加了管理难度；(3)共享性——多人可同时利用同一资源，容易引发公地悲剧；(4)波动性——资源量随环境条件和捕捞压力显著年际变化；(5)洄游性——许多种类具有跨行政区域的洄游习性，需要区域间协调管理。')

h3(doc, '5.4.2 捕捞过度的类型')
body(doc, '捕捞过度分为：(1)生物性捕捞过度——资源量下降到无法维持可持续产量水平，表现为CPUE持续下降、鱼体小型化低龄化；(2)经济性捕捞过度——捕捞的经济效益下降至低于成本水平。Russell公式B\u2082 = B\u2081 + R + G - M - Y是分析捕捞过度的理论基础：当Y > R + G - M时，资源量呈下降趋势，即出现捕捞过度。')

h3(doc, '5.4.3 主要管理制度')
body(doc, '我国主要的渔业资源管理制度包括：(1)禁渔期和禁渔区制度——海洋伏季休渔（每年5至9月）、长江十年禁渔（2021至2031年）等；(2)最小网目尺寸制度——中日渔业协定规定的网目标准以及我国统一标准（东黄海拖网网囊最小网目54mm，南海39mm）；(3)渔具限制制度——对拖网、刺网、定置网等渔具的类型、数量和规格进行限制；(4)捕捞许可证制度——实行捕捞许可证审批和渔船双控（船数和功率控制）制度；(5)水产种质资源保护区制度——划定重要水产种质资源的主要生长繁育区域进行特别保护。')

h2(doc, '5.5 渔业资源增殖保护费制度')
body(doc, f'渔业资源增殖保护费是根据{LQ}谁受益谁保护{RQ}原则，向从事渔业生产的单位和个人依法征收的专项费用。征收原则包括：取之于渔用之于渔、受益者负担、合理负担、分级征收。海洋渔业资源增殖保护费按渔船前3年采捕水产品年总产值的1%至3%征收。收入实行分级管理，沿海省级部门征收的90%留用地方，10%上缴海区渔政监督管理机构。')
body(doc, '渔业资源增殖保护费可用于：渔业资源增殖放流、资源调查评估、生态环境监测与保护、渔政执法管理、科学研究、宣传教育和受损渔民补偿等。')

h2(doc, '5.6 思考题')

questions = [
    ('名词解释：珍贵野生动物、濒危野生动物。',
     ['珍贵野生动物：在经济、科学、文化、教育、观赏等方面具有特殊重要价值，且数量稀少的野生动物物种。',
      '濒危野生动物：由于物种自身原因（繁殖力低、适应能力弱等）或受人类活动和自然灾害影响而有灭绝危险的野生动物物种。'],
     '《国家重点保护野生动物名录》是保护水生野生动物的核心法律依据，两类动物均受法律严格保护。'),
    ('简述保护水生野生动物的意义。',
     ['生态意义：水生野生动物是水域生态系统的关键组成部分，对维持生态平衡和食物网稳定具有不可替代的作用。',
      '科学意义：许多水生野生动物具有独特的生物学特性，如中华鲟的洄游习性、白鱀豚的回声定位能力，是生命科学和仿生学研究的重要对象。',
      '经济意义：水生野生动物具有观赏、旅游和潜在的经济开发价值。',
      '文化意义：水生野生动物是中华传统文化的重要组成部分。'],
     '保护水生野生动物就是保护水域生态系统的完整性和人类的共同遗产。'),
    ('水生野生动物保护与管理的基本制度有哪些？',
     ['特许捕捉制度——捕捉国家重点保护水生野生动物须经渔业行政主管部门批准。',
      '驯养繁殖许可证制度——驯养繁殖须取得许可证。',
      '经营利用管理制度——经营利用须遵守相关规定。',
      '运输管理制度——运输须持合法来源证明。',
      '进出口管理制度——须遵守CITES公约和国内法规。',
      '保护补偿制度——因保护造成的损失依法予以补偿。',
      '保护专项资金制度——设立专项资金用于保护工作。'],
     '七项制度构成了中国水生野生动物保护的法律制度体系。'),
    ('名词解释：渔业资源、捕捞过度、补充群体、剩余群体。',
     ['渔业资源：天然水域中具有开发利用价值的经济动植物种类和数量的总称。',
      '捕捞过度：捕捞强度超过资源再生能力，导致资源量持续下降的状态。分为生物性捕捞过度和经济性捕捞过度。',
      '补充群体：新补充到渔场中的个体，在生物学上指达到性成熟的个体，在渔业上指达到可捕规格的个体。',
      '剩余群体：上一周期未被捕获而存活下来的个体，与补充群体共同构成当前的渔业资源量。'],
     '四个概念是理解渔业资源动态和管理的基础术语。'),
    ('渔业资源有哪些自然特性？',
     ['可再生性——合理利用时能自我更新，需控制捕捞强度不超过资源再生能力。',
      '流动性——分布范围广且不断移动，需区域间协调管理。',
      '共享性——可被多个主体同时利用，需建立明确的产权制度和准入限制。',
      '波动性——资源量随环境和捕捞压力年际变化，管理需具有适应性。',
      '洄游性——许多种类有跨区域洄游习性，需国际或区域间合作管理。'],
     '渔业资源的自然特性决定了其管理不能照搬陆地资源管理模式，需建立适应性的综合管理制度。'),
    ('简述禁渔区和禁渔期的作用。',
     ['保护产卵亲鱼：在繁殖季节禁止捕捞，确保亲鱼顺利完成繁殖活动。',
      '保护幼鱼资源：为幼鱼提供生长的时间和空间，提高补充群体的数量和质量。',
      '促进资源恢复：降低捕捞压力，为衰退的资源提供恢复机会。长江十年禁渔是典型案例。',
      '维护生态平衡：保护水域生态系统的完整性和生物多样性。'],
     '禁渔制度是渔业资源养护最基础、最有效的管理措施之一。我国已建立了海洋伏季休渔和长江禁渔等多层次的禁渔制度体系。'),
    ('渔业资源增殖保护费征收的原则是什么？可用于哪些项目的开支？',
     ['征收原则：取之于渔用之于渔、受益者负担、合理负担、分级征收。',
      '收费标准：按产值的1%至5%不等，根据捕捞方式和作业水域差异确定差别费率。',
      '分级管理：90%留地方用于本地资源保护，10%上交中央用于全国性保护工作。',
      '可用于：增殖放流、调查评估、生态监测、渔政执法、科学研究、宣传教育、受损渔民补偿等。'],
     '渔业资源增殖保护费制度是落实谁受益谁保护原则的重要经济手段。'),
]

for i, (q_text, a_items, e_text) in enumerate(questions, 1):
    q(doc, i, q_text)
    ans(doc)
    for j, a in enumerate(a_items, 1):
        ap(doc, j, a)
    exp(doc, e_text)

doc.save(DOCX)
total=sum(len(p.text) for p in doc.paragraphs)
bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
print(f'文档已保存: {DOCX}')
print(f'总字符数: {total}')
print(f'粗体运行: {bold_c}')
print(f'表格数: {len(doc.tables)}')
