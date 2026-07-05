#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生物统计课件复习资料生成脚本 - 基于extracted_content.md
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OCR_DIR = os.path.expanduser("~/.hermes/project-data/生物统计/ocr")
CONTENT_PATH = os.path.expanduser("~/.hermes/project-data/生物统计/extracted_content_v3.md")
OUTPUT_PATH = os.path.expanduser("~/Downloads/学习/生物统计/生物统计课件整理.docx")

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_SIZE = Pt(10.5)

def set_run_font(run, font_cn=FONT_CN, font_en=FONT_EN, size=FONT_SIZE):
    run.font.size = size
    run.font.name = font_en
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    run.font.bold = False

def add_heading(doc, text, level=2):
    style_name = 'Heading 2' if level == 2 else 'Heading 3'
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    run = p.add_run(text)
    set_run_font(run, font_cn='黑体', size=Pt(14) if level==2 else Pt(12))
    run.font.bold = False
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = None

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.first_line_indent = Pt(21) if indent else None

def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(caption)
    set_run_font(run)

def create_three_line_table(doc, headers, rows, caption=None):
    if caption:
        add_table_caption(doc, caption)
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=Pt(9))
        p.alignment = 1
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=Pt(9))
            p.alignment = 1
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    tbl_borders.set(qn('w:top'), 'nil')
    tbl_borders.set(qn('w:left'), 'nil')
    tbl_borders.set(qn('w:bottom'), 'nil')
    tbl_borders.set(qn('w:right'), 'nil')
    tbl_borders.set(qn('w:insideH'), 'nil')
    tbl_borders.set(qn('w:insideV'), 'nil')
    tbl_pr.append(tbl_borders)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            if row_idx == 0:
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')
                tc_borders.append(top)
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                tc_borders.append(bottom)
            if row_idx == len(table.rows) - 1:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                tc_borders.append(bottom)
            tc_pr.append(tc_borders)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)

def add_qa(doc, questions):
    add_heading(doc, "本章习题", level=3)
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q['q']}")
        set_run_font(run)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(6)
        p = doc.add_paragraph()
        run = p.add_run("答案:")
        set_run_font(run)
        p.paragraph_format.first_line_indent = Pt(21)
        for j, point in enumerate(q['a'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"({j}) {point}")
            set_run_font(run)
            p.paragraph_format.left_indent = Pt(42)
            p.paragraph_format.first_line_indent = Pt(0)
        if q.get('x'):
            p = doc.add_paragraph()
            run = p.add_run(f"解析: {q['x']}")
            set_run_font(run)
            p.paragraph_format.first_line_indent = Pt(21)

CHAPTER_QA = {
    "第0章 绪论": [
        {"q": "生物统计学的主要内容有哪些?", "a": ["试验设计：合理安排试验，获取可靠数据", "统计分析：对数据进行整理、分析和推断"], "x": "生物统计学包括试验设计和统计分析两大主要内容。"}
    ],
    "第1章 试验资料的整理和特征数的计算": [
        {"q": "数量性状和质量性状的区别是什么?", "a": ["数量性状是可以用度量衡等计量工具测量的性状，如体长、体重等", "质量性状是只能观察而不能测量的性状，如颜色、性别等", "数量性状数据连续，质量性状数据间断"], "x": "正确区分性状类型是选择合适统计方法的基础。"},
        {"q": "算术平均数有哪些基本性质?", "a": ["离均差之和等于零:Σ(x- x̄)=0", "离均差平方和最小:Σ(x- x̄)²=最小值"], "x": "这两个性质是许多统计方法的数学基础。"}
    ],
    "第2章 理论分布与抽样分布": [
        {"q": "什么是小概率原理?", "a": ["概率很小的事件在一次试验中实际上几乎不可能发生", "它是显著性检验的基本依据", "一般认为小于0.05或0.01的概率为小概率"], "x": "小概率原理是假设检验的理论基础。"},
        {"q": "正态分布有哪些重要性质?", "a": ["曲线以μ为对称轴，左右对称", "在x=μ处取最大值", "曲线向两端无限延伸，以x轴为渐近线", "μ决定位置，σ决定形状"], "x": "正态分布是最重要的连续型分布。"}
    ],
    "第3章 统计推断": [
        {"q": "假设检验的基本步骤是什么?", "a": ["提出假设：建立无效假设H₀和备择假设Hₐ", "确定显著性水平α", "计算检验统计量", "做出统计推断"], "x": "假设检验是统计推断的核心方法。"},
        {"q": "什么是Ⅰ型错误和Ⅱ型错误?", "a": ["Ⅰ型错误：H₀为真却拒绝H₀，概率为α", "Ⅱ型错误：H₀为假却接受H₀，概率为β", "增大样本量可同时减小两类错误"], "x": "理解两类错误有助于正确解释检验结果。"}
    ],
    "第4章 卡方检验": [
        {"q": "χ²检验的适用条件是什么?", "a": ["用于计数资料或间断性数据的检验", "基本原理是比较实际观测值(O)与理论推算值(E)的偏离程度", "自由度df=1时需要进行连续性矫正"], "x": "χ²检验是处理分类数据的重要方法。"}
    ],
    "第5章 方差分析": [
        {"q": "方差分析的基本原理是什么?", "a": ["基于平方和的加和性原理", "将总平方和分解为处理平方和和误差平方和", "通过F检验判断处理因素效应的显著性"], "x": "方差分析用于比较多个平均数。"},
        {"q": "方差分析有哪些基本假设?", "a": ["各处理的观测值相互独立", "各处理的观测值服从正态分布", "各处理的方差相等(方差齐性)"], "x": "满足这些假设才能保证方差分析结果的可靠性。"}
    ],
    "第6章 回归与相关分析": [
        {"q": "回归分析和相关分析的区别是什么?", "a": ["回归分析具有预测性，相关分析不具备", "回归分析变量地位不平等，相关分析变量地位平等", "回归分析因变量是随机变量，相关分析所有变量都是随机变量"], "x": "回归和相关是两种不同的统计分析方法。"},
        {"q": "相关系数的性质有哪些?", "a": ["取值范围：-1≤r≤1", "r>0为正相关，r<0为负相关", "|r|越大表示相关越密切"], "x": "相关系数是衡量两个变量线性相关程度的指标。"}
    ]
}

def parse_md_to_doc(doc, md_text):
    """将markdown文本转换为Word文档段落"""
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 跳过分隔线
        if line == '---':
            i += 1
            continue
        
        # 二级标题 ##
        if line.startswith('## '):
            text = line[3:].strip()
            add_body(doc, text)
            i += 1
            continue
        
        # 三级标题 ###
        if line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, level=3)
            i += 1
            continue
        
        # 四级标题 ####
        if line.startswith('#### '):
            text = line[5:].strip()
            # 作为加粗段落处理
            add_body(doc, text)
            i += 1
            continue
        
        # 表格
        if '|' in line and i+1 < len(lines) and lines[i+1].strip().startswith('|---'):
            # 解析表格
            headers = [c.strip() for c in line.split('|') if c.strip()]
            i += 2  # 跳过分隔行
            rows = []
            while i < len(lines) and '|' in lines[i]:
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                if cells:
                    # 确保列数一致
                    while len(cells) < len(headers):
                        cells.append('')
                    cells = cells[:len(headers)]
                    rows.append(cells)
                i += 1
            if headers and rows:
                create_three_line_table(doc, headers, rows)
            continue
        
        # 列表项 -
        if line.startswith('- '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            add_body(doc, f"• {text}")
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+\.', line):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            add_body(doc, text)
            i += 1
            continue
        
        # 引用块 >
        if line.startswith('> '):
            text = line[2:].strip()
            add_body(doc, text)
            i += 1
            continue
        
        # 普通段落
        # 收集连续的非空行作为一个段落
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or next_line.startswith('#') or next_line.startswith('- ') or next_line.startswith('>') or next_line.startswith('|') or next_line == '---':
                break
            para_lines.append(next_line)
            i += 1
        
        text = ' '.join(para_lines)
        # 清理markdown格式标记
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 移除加粗
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # 移除斜体
        text = re.sub(r'`(.+?)`', r'\1', text)  # 移除代码标记
        add_body(doc, text)

def main():
    # 读取extracted_content.md
    with open(CONTENT_PATH, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 按章节分割
    sections = re.split(r'^# (第.+章.+)$', content, flags=re.MULTILINE)
    
    chapters = {}
    i = 1
    while i < len(sections):
        title = sections[i].strip()
        body = sections[i+1].strip() if i+1 < len(sections) else ""
        chapters[title] = body
        i += 2
    
    # 创建文档
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    style.font.size = FONT_SIZE
    
    # 主标题
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("水产生物统计 课件复习资料")
    set_run_font(run, font_cn='黑体', size=Pt(16))
    run.font.bold = False
    
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("华中农业大学水产学院")
    set_run_font(run, size=Pt(12))
    
    add_body(doc, "本资料根据课程PPT整理，涵盖各章核心知识点、重要公式和典型例题，供复习参考。")
    doc.add_page_break()
    
    # 生成各章内容
    chapter_order = [
        "第0章 绪论",
        "第1章 试验资料的整理和特征数的计算",
        "第2章 理论分布与抽样分布",
        "第3章 统计推断",
        "第4章 卡方检验",
        "第5章 方差分析",
        "第6章 回归与相关分析"
    ]
    
    for ch_title in chapter_order:
        if ch_title in chapters:
            # 提取章节号和标题
            match = re.match(r'第(.+?)章\s*(.*)', ch_title)
            if match:
                ch_num = match.group(1)
                ch_name = match.group(2)
                heading_text = f"{ch_num}. {ch_name}" if ch_num != "0" else f"0. {ch_name}"
            else:
                heading_text = ch_title
            
            add_heading(doc, heading_text, level=2)
            parse_md_to_doc(doc, chapters[ch_title])
            # 添加Q&A
            if ch_title in CHAPTER_QA:
                add_qa(doc, CHAPTER_QA[ch_title])
            doc.add_page_break()
    
    # 附录：公式汇总
    if '附录' in content:
        # 找到附录内容
        appendix_start = content.find('# 附录')
        if appendix_start >= 0:
            appendix_content = content[appendix_start:]
            add_heading(doc, "附录：核心公式汇总", level=2)
            parse_md_to_doc(doc, appendix_content)
    
    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"文档已生成: {OUTPUT_PATH}")
    print(f"总字符数: {total_chars}")
    print(f"表格数: {len(doc.tables)}")
    print(f"加粗: {sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)}")

if __name__ == "__main__":
    main()
