from docx import Document

doc = Document(r'C:\Users\Rei\Downloads\学习\生物统计\生物统计课件整理.docx')
print(f'总段落数: {len(doc.paragraphs)}')
print(f'总表格数: {len(doc.tables)}')
print()

# 前30段看结构
for i, p in enumerate(doc.paragraphs[:30]):
    style = p.style.name if p.style else 'None'
    fi = p.paragraph_format.first_line_indent
    left = p.paragraph_format.left_indent
    text = p.text[:100] if p.text else '(empty)'
    print(f'[{i:3d}] style={style:15s} fi={fi} left={left} | {text}')

print('...\n...\n')

# 搜寻章节标题
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith(('第', '章', '0.', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')):
        print(f'[{i:3d}] {p.text[:80]}')
    elif 'Heading' in (p.style.name if p.style else ''):
        print(f'[{i:3d}] [{p.style.name:12s}] {p.text[:80]}')
