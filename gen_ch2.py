#!/usr/bin/env python3
"""Append Chapter 2 to existing docx."""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")
LQ='\u201c'; RQ='\u201d'

def sf(r, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False, italic=False):
    r.font.size=sz; r.font.name=en
    r.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.bold=bold; r.font.italic=italic

def body(doc, txt, indent=True):
    p=doc.add_paragraph(); r=p.add_run(txt); sf(r)
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def h1(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 1']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(16),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

def h2(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 2']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(14),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def h3(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 3']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(12),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def q(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'{n}. {txt}'); sf(r,sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def ans(doc):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run('答案:'); sf(r)

def ap(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'({n}) {txt}'); sf(r,sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210'); pPr.append(ind)

def exp(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run(f'解析: {txt}'); sf(r,sz=Pt(10))

doc = Document(DOCX)
doc.add_page_break()

h1(doc, '2. 鱼类资源的生物学基础')

h2(doc, '2.1 鱼类种群的特征与结构')
h3(doc, '2.1.1 种群的基本概念')
body(doc, '种群（Population）是指一定空间范围内同种生物个体的集合，具有一定的分布范围，在空间、数量和遗传上具有相对独立性。种群是渔业资源管理的基本单元，正确划分种群边界是资源评估和管理的基础。')
body(doc, '种群的形成主要基于两种隔离机制：(1)地理隔离——海洋、山脉、河流等地理障碍导致种群间无法自由交配，长期的地理隔离促进种群间的遗传分化。不同河流中的同一鱼类往往形成不同的地理种群。(2)生殖隔离——由于生殖期不同（如产卵季节差异）、生殖行为差异或生殖器官结构差异等形成的隔离，即使在同一地理区域也能保持种群间的独立性。')
body(doc, '种群的特征包括三个方面：形态学特征（体形、体色、鳞片、鳍条、鳃耙数、脊椎骨数等可量可数性状）、生物学特征（生长率、成熟年龄、繁殖力、寿命等）和生理生化特征（同工酶谱、DNA序列等）。种群鉴别的主要方法有统计学方法（形态测量多元分析）、生化遗传方法（同工酶电泳）和分子生物学方法（线粒体DNA标记、微卫星标记等）。')

h3(doc, '2.1.2 种群结构')
body(doc, '鱼类种群结构主要包括年龄组成、长度与重量组成和性比组成三个方面。经济鱼类的寿命多在2龄至几十龄之间，年龄结构可分为单龄群体结构（年龄组只有一个，如一年生鱼类）和多年龄群体结构两种类型。')
body(doc, '长度与重量组成的资料比年龄资料更容易获得，可迅速给出百分组成或直方图供分析使用。营养条件是影响鱼类生长和体长体重关系的重要因素。性比组成方面：多数鱼类生殖群体的雌雄比例接近1:1；筑巢产卵鱼类（如乌鳢）产卵时雌雄比例约1:1；某些特殊种类如异育银鲫（中科3号）生殖群体中雌体占优势；部分鱼类存在雌雄同体性转化现象，如石斑鱼在生活史中具有由雌变雄的性反转过程。')

h3(doc, '2.1.3 种群数量与估算')
body(doc, '决定种群数量的因素包括：生物潜能（Biotic potential）——适合种群延续与增加其数量的潜在能力；繁殖力——单位时间内种群产生新个体的能力；抗逆性——种群抵抗不利环境维持自我发展的能力；种间关系——Gause原理（竞争排斥原理）指出两个生态位相同的物种不能长期共存。')
body(doc, '种群数量估算方法有直接计数法、抽样法和标志重捕法。标志重捕法（Lincoln-Petersen法）的基本公式为：N = M \u00d7 n / m，其中M为标记并释放的个体数，n为第二次捕获的总个体数，m为第二次捕获中带标记的个体数。该方法假设标记个体与未标记个体混合均匀，且标记不影响存活和行为。')

h2(doc, '2.2 鱼类的年龄与生长')
body(doc, '鱼类年龄与生长研究具有重要的渔业意义：(1)为制定合理捕捞强度提供科学依据；(2)确定合理捕捞规格；(3)为渔情预报提供基础资料；(4)拟定水域养殖种类及措施；(5)提高种类引种和驯化效果；(6)鱼类生长特点是研究种群特征的重要依据。')
body(doc, '年龄鉴定主要依据鳞片、耳石、脊椎骨和鳃盖骨等硬组织上的年轮特征。鳞片年轮清晰、易于采集，是最常用的年龄鉴定材料；耳石的微化学分析可追溯鱼类的迁徙历史和生活史事件。von Bertalanffy生长方程（VBGF）是描述鱼类生长最常用的数学模型：L\u209c = L\u221e(1 - e\u207b\u1d3e\u1d49\u02e2\u207b\u1d49\u2070），其中L\u209c为t龄时的体长，L\u221e为理论渐近体长，K为生长系数（曲率参数），t\u2080为理论生长起点年龄（体长为0时的理论年龄）。')

h2(doc, '2.3 鱼类性成熟、繁殖习性与繁殖力')
body(doc, '鱼类性成熟受遗传因素和环境条件的共同影响。同种鱼类的性成熟年龄在不同纬度、不同水温条件下可能有显著差异。性腺成熟度通常用性腺成熟系数表示：GSI = 性腺重/体重 \u00d7 100%。繁殖力（Fecundity）指雌鱼在繁殖季节所产卵子的数量，是评估种群生产潜力的重要参数。不同鱼类繁殖力差异很大，鲟鱼一次可产数百万粒卵，而某些卵胎生鱼类仅产数十尾幼鱼。繁殖力受鱼体大小、年龄、营养状况和环境条件等多种因素影响。')

h2(doc, '2.4 鱼类的集群与洄游')
h3(doc, '2.4.1 集群')
body(doc, '鱼类集群（Schooling）是重要的社会行为，具有降低被捕食风险、提高觅食效率、减少游泳能耗和促进生殖活动等生物学意义。集群行为的研究方法包括水下观察、声学探测和视频分析等。')

h3(doc, '2.4.2 洄游')
body(doc, '鱼类洄游（Migration）是周期性、定向性的长距离移动行为。按类型分为：(1)溯河洄游（Anadromous）——在海洋中生长，性成熟后上溯到河流中繁殖，如鲑科鱼类（太平洋鲑的洄游路线是经典案例）；(2)降海洄游（Catadromous）——在淡水中生长，性成熟后降河入海繁殖，如鳗鲡；(3)淡水鱼类洄游（Potamodromous）——在淡水水域内进行的洄游，如四大家鱼的产卵洄游。')
body(doc, '鱼类洄游的机制涉及环境因子（水温、盐度、水流、光照等）和内在因素（生理状态、激素水平等）的综合作用。定向机制包括太阳罗盘定向、地磁定向、嗅觉定向和记忆定向等多种方式。洄游通道的畅通对鱼类完成生活史至关重要，水利工程建设对洄游鱼类的负面影响一直是渔业生态学关注的重点问题。')

h2(doc, '2.5 思考题')

q(doc, 1, '解释种群(Population)、生物潜能(Biotic potential)、环境容纳量(Carrying capacity)、内禀增长率(Intrinsic rate)的概念。')
ans(doc)
ap(doc, 1, '种群（Population）：一定空间范围内同种生物个体的集合，在空间、数量和遗传上具有相对独立性，是渔业资源管理的基本单元。')
ap(doc, 2, '生物潜能（Biotic potential）：适合种群延续与增加其数量的潜在能力，即在环境允许的前提下种群所能达到的最佳数量。')
ap(doc, 3, '环境容纳量（Carrying capacity）：在特定环境条件下，一个生态系统所能维持的最大种群数量，通常用K表示。')
ap(doc, 4, '内禀增长率（Intrinsic rate）：在最适环境条件下（食物充足、空间充裕、无敌害），种群所能达到的最大瞬时增长率，通常用r表示。')
exp(doc, '四个概念相互关联：生物潜能决定种群增长潜力，环境容纳量设定增长上限，内禀增长率反映理论增长速率。在渔业管理中，常通过r和K确定最大持续产量MSY = rK/4。')

q(doc, 2, '种群的形成原因有哪些？鉴别种群的技术方法有哪些？')
ans(doc)
ap(doc, 1, '种群形成原因：(1)地理隔离——海洋、山脉、河流等地理障碍导致种群间无法自由交配和基因交流；(2)生殖隔离——生殖期不同、生殖行为差异或生殖器官结构差异导致的隔离。')
ap(doc, 2, '鉴别方法：(1)统计学方法——形态测量、多元分析；(2)生化遗传方法——同工酶电泳分析；(3)分子生物学方法——线粒体DNA标记、微卫星标记、SNP分析等。')
exp(doc, '种群鉴别是渔业资源管理的基础工作，准确的种群界定对制定分区域捕捞配额和保护措施至关重要。')

q(doc, 3, '简述鱼类年龄与生长研究在渔业上的意义。')
ans(doc)
ap(doc, 1, '为制定合理的捕捞强度提供科学依据——通过生长参数确定最适捕捞规格和捕捞强度。')
ap(doc, 2, '确定合理的捕捞规格——根据生长模型计算首次捕捞年龄t\u2091，使单位补充量渔获量最大化。')
ap(doc, 3, '为渔情预报提供基础资料——年龄组成和生长数据是资源评估模型的重要输入参数。')
ap(doc, 4, '拟定水域养殖种类及措施——根据生长特性选择适宜养殖种类和制定养殖管理措施。')
ap(doc, 5, '提高引种和驯化效果——了解目标种的生长规律有助于制定科学引种方案。')
ap(doc, 6, '鱼类生长特点是研究种群特征的重要依据——生长差异反映环境质量和种群密度变化。')
exp(doc, 'von Bertalanffy生长方程L\u209c = L\u221e(1 - e\u207b\u1d3e\u02e2\u207b\u1d49\u2070）是描述鱼类生长最常用的数学模型，广泛应用于资源评估。')

q(doc, 4, '鱼类的洄游有哪些类型？各有什么特点和代表性种类？')
ans(doc)
ap(doc, 1, '溯河洄游：在海洋中生长，性成熟后上溯到河流中繁殖。代表性种类有鲑科鱼类、中华鲟、鲥鱼等。特点：洄游距离远、方向性强，产卵后部分种类死亡。')
ap(doc, 2, '降海洄游：在淡水中生长，性成熟后降河入海繁殖。代表性种类有鳗鲡（日本鳗鲡）。特点：成体在淡水生长多年后一次性降海产卵，产卵后死亡。')
ap(doc, 3, '淡水鱼类洄游：在淡水水域内进行。代表性种类有四大家鱼（青鱼、草鱼、鲢、鳙）的产卵洄游和青海湖裸鲤的生殖洄游。特点：洄游距离相对较短，一般在同一水系内完成。')
exp(doc, '鱼类洄游机制涉及环境因子（水温、盐度、水流、光照等）和内在因素的综合作用。定向机制包括太阳罗盘、地磁、嗅觉和记忆定向等。洄游通道畅通对鱼类完成生活史至关重要。')

doc.save(DOCX)
total=sum(len(p.text) for p in doc.paragraphs)
bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
print(f'文档已保存: {DOCX}')
print(f'总字符数: {total}')
print(f'粗体运行: {bold_c}')
print(f'表格数: {len(doc.tables)}')
