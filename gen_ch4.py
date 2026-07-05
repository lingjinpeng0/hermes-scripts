#!/usr/bin/env python3
"""Append Chapter 4 to existing docx."""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")

def sf(r, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False, italic=False):
    r.font.size=sz; r.font.name=en
    r.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.bold=bold; r.font.italic=italic

def body(doc, txt, indent=True):
    p=doc.add_paragraph(); r=p.add_run(txt); sf(r)
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def h1(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 1']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(16),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

def h2(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 2']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(14),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def h3(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 3']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(12),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def tbl(doc, hd, rows, cap=None):
    if cap:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(cap); sf(r); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.first_line_indent=None
    t=doc.add_table(rows=1+len(rows), cols=len(hd))
    t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tl=t._tbl; tp=tl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); tl.insert(0,tp)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct'); tp.append(tw)
    tb=OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'nil'); tb.append(b)
    tp.append(tb)
    for ci,h in enumerate(hd):
        c=t.rows[0].cells[ci]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Pt(0)
        r=p.add_run(h); sf(r,sz=Pt(10))
        _bdr(c,'12','6')
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if len(str(v))<30 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent=Pt(0)
            r=p.add_run(str(v)); sf(r,sz=Pt(9))
    for ci in range(len(hd)):
        c=t.rows[-1].cells[ci]; tcp=c._tc.get_or_add_tcPr()
        tcb=OxmlElement('w:tcBorders'); bt=OxmlElement('w:bottom')
        bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'12')
        bt.set(qn('w:space'),'0'); bt.set(qn('w:color'),'000000')
        tcb.append(bt); tcp.append(tcb)
    for row in tl.findall(qn('w:tr')):
        trp=row.find(qn('w:trPr'))
        if trp is None: trp=OxmlElement('w:trPr'); row.insert(0,trp)
        old=trp.find(qn('w:cantSplit'))
        if old is not None: trp.remove(old)
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trp.append(cs)
    tPr=tl.tblPr
    tjc=OxmlElement('w:jc'); tjc.set(qn('w:val'),'center'); tPr.append(tjc)

def _bdr(c,ts='12',bs='6'):
    tcp=c._tc.get_or_add_tcPr(); tcb=OxmlElement('w:tcBorders')
    for s,sz in [('top',ts),('bottom',bs)]:
        e=OxmlElement(f'w:{s}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); tcb.append(e)
    tcp.append(tcb)

def q(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'{n}. {txt}'); sf(r,sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def ans(doc):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run('答案:'); sf(r)

def ap(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'({n}) {txt}'); sf(r,sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210'); pPr.append(ind)

def exp(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run(f'解析: {txt}'); sf(r,sz=Pt(10))

doc = Document(DOCX)
doc.add_page_break()

h1(doc, '4. 渔业资源评估')

h2(doc, '4.1 动态综合模型')
h3(doc, '4.1.1 概述')
body(doc, '动态综合模型（Dynamic Pool Model）又称单位补充量渔获量模型（Yield per Recruit Model），其基本思路是将种群资源量分解为单位补充量的变化过程。模型不直接研究补充量本身的变化，而是分析生长、死亡和捕捞等因素对单位补充量渔获量（Y/R）的影响，通过优化捕捞死亡系数F和首次捕捞年龄t\u2091实现渔获量的最大化。')
body(doc, '模型的两个基本特征：(1)将补充量R视为独立变量，不建立补充量与亲体量之间的函数关系；(2)关注Y/R随F和t\u2091的变化规律，通过等产量曲线分析最优组合。')

h3(doc, '4.1.2 B-H模型')
body(doc, 'B-H模型（Beverton-Holt模型）的假设条件包括：补充量恒定不变；自然死亡系数M为常数；生长参数（L\u221e\u3001K\u3001t\u2080）保持稳定；捕捞死亡系数F为常数；完全补充后渔具选择性一致。模型需要已知的参数包括：生长参数L\u221e\u3001K\u3001t\u2080，自然死亡系数M，首次捕捞年龄t\u2091，捕捞死亡系数F，极限年龄t\u03bb。')
body(doc, 'B-H模型的核心方程为：Y/R = F \u00d7 \u2211W\u209c \u00d7 e\u207b\u1d39\u02e3\u207b\u1d39\u02e3，通过对各个年龄组的体重和存活率求和，计算单位补充量的总渔获量。von Bertalanffy生长方程用于计算各年龄组的平均体重，指数死亡模型用于计算各年龄组的存活率。')

h3(doc, '4.1.3 F和t\u2091对渔获量的影响')
body(doc, 'F（捕捞死亡系数）的影响：在固定的t\u2091下，Y/R随F的增加先增加后减少，存在一个最适捕捞死亡系数F\u2092pt。当F小于F\u2092pt时，增加F可以提高渔获量；当F大于F\u2092pt后，继续增加F会导致Y/R下降，即出现生长性捕捞过度。')
body(doc, 't\u2091（首次捕捞年龄）的影响：在固定的F下，Y/R随t\u2091增大先增后减，存在最适首次捕捞年龄。t\u2091过小（捕捞幼鱼）导致大多数个体在达到最佳生长体重前被捕捞，造成生长性捕捞过度；t\u2091过大会使更多的个体因自然死亡而损失，渔业效益不佳。F和t\u2091之间存在交互作用，需要通过等产量曲线分析最优组合。')

h2(doc, '4.2 剩余产量模型')
body(doc, '剩余产量模型（Surplus Production Model）将种群视为一个整体，不需要详细的年龄结构数据，适用于数据有限的渔业。Schaefer模型是最经典的形式：dB/dt = rB(1 - B/K) - qfB，其中r为内禀增长率，K为环境容纳量，q为捕捞系数，f为捕捞努力量。')
body(doc, '在平衡状态下，渔获量Y = qfB*，通过对dB/dt = 0求解可得B* = K(1 - qf/r)。代入可得平衡渔获量Y = qfK(1 - qf/r)，进一步求导可得：最大持续产量MSY = rK/4，最适捕捞努力量f\u2098SY = r/(2q)，最适资源量B\u2098SY = K/2。参数估算通常通过CPUE与捕捞努力量的回归分析实现。')

h2(doc, '4.3 亲体与补充量关系模型')
body(doc, 'Ricker模型：R = \u03b1Se\u207b\u03b2S，适用于补充量在亲体量中等时达到最大、之后呈下降趋势的情形，典型例子为鲑科鱼类。Beverton-Holt繁殖模型：R = 1/(\u03b1 + \u03b2/S)，适用于补充量随亲体量增加而趋于渐近值的情形，大多数海洋鱼类符合此模型。此外还有Cushing模型（幂函数形式）、Shepherd模型和Gamma模型等。')

h2(doc, '4.4 资源量估算方法')

tbl(doc,
    ['估算方法', '所需数据', '优点', '局限性'],
    [
        ['扫海面积法', '拖网采样数据', '直接可靠', '耗时费力'],
        ['标志放流法', '标记重捕数据', '适用范围广', '标记可能影响行为'],
        ['Leslie/Delury法', 'CPUE+努力量', '计算简单', '需捕捞死亡率高'],
        ['水声学法', '声学探测数据', '大面积快速', '种类鉴别困难'],
        ['鱼卵仔鱼法', '鱼卵仔鱼密度', '不伤成鱼', '仅适用产卵期'],
    ],
    cap='表4-1.主要资源量估算方法比较')

body(doc, '常用的渔业资源量估算方法包括：(1)扫海面积法——利用拖网采样，根据扫海面积（拖曳距离\u00d7网口宽度\u00d7捕获效率）和渔获密度推算资源量；(2)鱼卵仔鱼法——通过调查鱼卵和仔鱼密度估算亲鱼资源量；(3)标志放流法——包括Lincoln法、Schnabel法和Jolly-Seber法；(4)Leslie法和Delury法——利用CPUE随累积捕捞努力量的下降趋势推算初始资源量；(5)营养动态法——根据生态系统能流推算渔业资源生产潜力；(6)水声学法——利用声学探测设备对鱼群进行定量评估。')

h2(doc, '4.5 思考题')

q(doc, 1, 'B-H模型中捕捞死亡系数F和首次捕捞年龄t\u2091分别对渔获量有何影响？两者如何交互？')
ans(doc)
ap(doc, 1, 'F的影响：在t\u2091固定时，Y/R随F增加先增后减，存在最适F\u2092pt。F \u003c F\u2092pt时增加F可提高产量；F \u003e F\u2092pt时导致捕捞过度，Y/R下降。')
ap(doc, 2, 't\u2091的影响：在F固定时，Y/R随t\u2091增大先增后减，存在最适首次捕捞年龄。t\u2091过小造成生长性捕捞过度；t\u2091过大则自然死亡消耗过多。')
ap(doc, 3, '交互作用：F和t\u2091相互影响，需在二维平面上绘制等产量曲线确定最优组合。当F较高时，适当提高t\u2091可补偿过度捕捞的影响；当t\u2091较小时，需降低F避免生长性捕捞过度。')
exp(doc, '合理的渔业管理需要同时优化捕捞强度和最小可捕规格，两者必须统筹考虑。')

q(doc, 2, '简述剩余产量模型中Schaefer模型的数学形式及其在渔业管理中的应用。')
ans(doc)
ap(doc, 1, 'Schaefer模型：dB/dt = rB(1 - B/K) - qfB。模型假设种群增长符合逻辑斯谛方程。')
ap(doc, 2, 'MSY = rK/4，f\u2098SY = r/(2q)，B\u2098SY = K/2。这些参数是确定捕捞限额和捕捞强度控制目标的理论依据。')
ap(doc, 3, '管理应用：确定最大持续产量MSY作为捕捞限额上限参考值；计算最适捕捞努力量f\u2098SY作为控制目标；通过当前CPUE与历史数据对比判断资源利用状态。')
exp(doc, '剩余产量模型不需要详细的年龄结构数据，适用于数据有限的发展中国家渔业。但模型假设（平衡状态、环境稳定等）在实际中往往难以完全满足。')

q(doc, 3, '常用的渔业资源量估算方法有哪些？各有什么优缺点？')
ans(doc)
ap(doc, 1, '扫海面积法：利用拖网采样直接推算资源量，结果直接可靠但耗时费力。')
ap(doc, 2, '标志放流法：通过标记重捕数据估算，适用范围广但标记可能影响鱼类行为。')
ap(doc, 3, 'Leslie/Delury法：利用CPUE下降趋势推算，计算简单但需要较高的捕捞死亡率。')
ap(doc, 4, '水声学法：利用声学探测进行大面积快速评估，但种类鉴别困难。')
exp(doc, '不同方法适用于不同的渔业类型和数据条件，实际应用中常采用多种方法交叉验证以提高评估准确性。')

doc.save(DOCX)
total=sum(len(p.text) for p in doc.paragraphs)
bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
print(f'文档已保存: {DOCX}')
print(f'总字符数: {total}')
print(f'粗体运行: {bold_c}')
print(f'表格数: {len(doc.tables)}')
