#!/usr/bin/env python3
"""Append Chapter 3 to existing docx."""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")
LQ='\u201c'; RQ='\u201d'

def sf(r, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False, italic=False):
    r.font.size=sz; r.font.name=en
    r.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.bold=bold; r.font.italic=italic

def body(doc, txt, indent=True):
    p=doc.add_paragraph(); r=p.add_run(txt); sf(r)
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def h1(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 1']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(16),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

def h2(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 2']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(14),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def h3(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 3']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(12),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def tbl(doc, hd, rows, cap=None):
    if cap:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(cap); sf(r); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.first_line_indent=None
    t=doc.add_table(rows=1+len(rows), cols=len(hd))
    t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tl=t._tbl; tp=tl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); tl.insert(0,tp)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct'); tp.append(tw)
    tb=OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'nil'); tb.append(b)
    tp.append(tb)
    for ci,h in enumerate(hd):
        c=t.rows[0].cells[ci]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Pt(0)
        r=p.add_run(h); sf(r,sz=Pt(10))
        _bdr(c,'12','6')
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if len(str(v))<30 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent=Pt(0)
            r=p.add_run(str(v)); sf(r,sz=Pt(9))
    for ci in range(len(hd)):
        c=t.rows[-1].cells[ci]; tcp=c._tc.get_or_add_tcPr()
        tcb=OxmlElement('w:tcBorders'); bt=OxmlElement('w:bottom')
        bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'12')
        bt.set(qn('w:space'),'0'); bt.set(qn('w:color'),'000000')
        tcb.append(bt); tcp.append(tcb)
    for row in tl.findall(qn('w:tr')):
        trp=row.find(qn('w:trPr'))
        if trp is None: trp=OxmlElement('w:trPr'); row.insert(0,trp)
        old=trp.find(qn('w:cantSplit'))
        if old is not None: trp.remove(old)
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trp.append(cs)
    tPr=tl.tblPr
    tjc=OxmlElement('w:jc'); tjc.set(qn('w:val'),'center'); tPr.append(tjc)

def _bdr(c,ts='12',bs='6'):
    tcp=c._tc.get_or_add_tcPr(); tcb=OxmlElement('w:tcBorders')
    for s,sz in [('top',ts),('bottom',bs)]:
        e=OxmlElement(f'w:{s}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); tcb.append(e)
    tcp.append(tcb)

def q(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'{n}. {txt}'); sf(r,sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def ans(doc):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run('答案:'); sf(r)

def ap(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'({n}) {txt}'); sf(r,sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210'); pPr.append(ind)

def exp(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run(f'解析: {txt}'); sf(r,sz=Pt(10))

doc = Document(DOCX)
doc.add_page_break()

h1(doc, '3. 鱼类资源数量的变动')

h2(doc, '3.1 研究渔业资源数量变动的基本单位')
body(doc, '渔业资源数量变动研究涉及的基本概念包括：资源总量（某一水域中某种或某类群渔业生物的总生物量）、可捕资源量（超过最小可捕规格、可供捕捞的部分）、预备资源量（即将进入可捕规格的资源量）。')
body(doc, '基本单位从大到小包括：(1)种群（种族）——具有一定形态特征、生理特性和生态要求的独立繁殖群体，是渔业管理的基本单元；(2)亚种群或种下群——种群内具有局部繁殖习性和分布特点的较小群体，在遗传上与种群其他部分存在一定差异；(3)群体、资源群体——根据特定管理目标划分的利用单元。')

h2(doc, '3.2 渔业资源数量变动的基本原因')
body(doc, '渔业资源数量变动受四类因素的综合影响：(1)生物学特征——繁殖能力、生长速率、自然死亡等种群内在参数，其中补充量（Recruitment）是新个体加入种群的过程，其大小直接影响资源量的变化趋势；(2)环境因素——水温、盐度、溶解氧、营养盐等理化环境条件，水温影响代谢速率和性腺发育，溶氧量影响摄食效率和活动能力；')
body(doc, '(3)生物因素——饵料生物的丰歉直接影响鱼类生长和肥满度；种间竞争（Gause原理）决定生态位的分化与重叠；敌害生物影响幼鱼的存活率；(4)人为因素——捕捞是最主要的人为影响因素，过度捕捞导致亲鱼减少、补充量下降；水利工程阻断洄游通道；水体污染破坏栖息地质量。')

h2(doc, '3.3 渔业资源数量变动的基本模型')
body(doc, 'Russell模型（1932）是渔业资源数量变动的基础分析框架：B\u2082 = B\u2081 + R + G - M - Y。其中B\u2082和B\u2081分别为期末和期初资源量，R为补充量，G为生长增量，M为自然死亡量，Y为渔获量。当Y < R + G - M时资源量增加，当Y > R + G - M时资源量减少，即出现捕捞过度。')
body(doc, 'Schaefer和Beverton用相对变化率替代绝对数量，发展了两大类模型：(1)剩余产量模型（Surplus Production Model）——将种群视为整体，研究生物量变化，核心为S型增长曲线和最大持续产量MSY = rK/4；(2)动态综合模型（Dynamic Pool Model）——考虑年龄结构，研究单位补充量渔获量Y/R随F和t\u2091的变化。')
body(doc, '补充群体（Recruitment group）在生物学上指达到性成熟并参与繁殖的个体，在渔业上指达到可捕规格并进入渔场的个体。两者在时间上通常不一致——鱼类可能在达到性成熟后还需一段时间生长才能达到可捕规格。剩余群体（Remaining group）指上一周期未被捕获而存活下来的个体。')

h2(doc, '3.4 捕捞对渔业资源和渔获量的影响')
body(doc, '捕捞努力量（Fishing effort）的标准化是资源评估的基础工作。CPUE（单位捕捞努力量渔获量）是资源量变动的重要指标，通常假设CPUE与资源量成正比。标准化方法包括效能比法和GLM（广义线性模型）法。')
body(doc, '渔具选择性（Gear selectivity）反映了不同渔具对不同大小鱼类的捕捞效率差异。拖网的选择性相对平缓——体长稍大于网目即被保留；刺网的选择性呈钟形曲线——特定体长范围的鱼最易被捕捞。选择性参数包括L\u2080\u2085（50%被保留的体长）和选择系数。了解渔具选择性对制定最小网目尺寸、保护幼鱼资源具有重要意义。')

tbl(doc,
    ['渔具类型', '选择性特征', '主要捕捞对象'],
    [
        ['拖网(Trawling)', '选择性较平缓', '底层和近底层鱼类'],
        ['流刺网(Gillnetting)', '钟形选择曲线', '鲐、鳓、鲳、梭子蟹等'],
        ['围网(Purse Seine)', '选择性较高', '中上层集群鱼类'],
        ['延绳钓(Longline)', '选择性较强', '金枪鱼、鲨鱼等'],
        ['定置网(Stow Net)', '选择性较低', '多种沿岸鱼类'],
    ],
    cap='表3-1.主要渔具类型及其选择性特征')

h2(doc, '3.5 思考题')

q(doc, 1, '补充群体在生物学与渔业上的含义有何不同？请结合具体案例说明。')
ans(doc)
ap(doc, 1, '生物学含义：指达到性成熟、首次参与繁殖的个体。从生物学角度完成了性腺发育并具备繁殖能力，是种群延续的关键组分。')
ap(doc, 2, '渔业含义：指达到可捕规格（超过最小可捕体长或年龄）并进入渔场的个体。在渔业上具有捕捞价值，进入捕捞群体的范围。')
ap(doc, 3, '差异与应用：生物学补充与渔业补充在时间上通常不一致——鱼类可能在达到性成熟后还需一段时间生长才能达到可捕规格。合理的管理策略应确保个体至少有一次繁殖机会后再被捕捞，即最小可捕体长应设定在性成熟体长以上。')
exp(doc, '正确理解补充群体的两种含义是确定最小网目尺寸和开捕规格的理论基础，对渔业可持续发展至关重要。')

q(doc, 2, 'Russell模型的基本表达式是什么？各参数的含义以及如何利用该模型判断资源量的变化趋势？')
ans(doc)
ap(doc, 1, '公式：B\u2082 = B\u2081 + R + G - M - Y')
ap(doc, 2, '参数含义：B\u2082为期未资源量；B\u2081为期初资源量；R为补充量（新加入个体的数量）；G为生长增量（个体生物量的增加总和）；M为自然死亡量（疾病、被捕食、衰老等非捕捞因素导致的死亡）；Y为渔获量（捕捞活动导致的死亡量）。')
ap(doc, 3, '判断方法：当R + G - M > Y时，B\u2082 > B\u2081，资源量增长；当R + G - M < Y时，B\u2082 < B\u2081，资源量衰退，即出现捕捞过度。')
exp(doc, 'Russell模型是所有渔业资源变动模型的基础框架，后来的剩余产量模型和动态综合模型都是在其基础上发展深化而来。')

q(doc, 3, '影响渔业资源数量变动的主要因素有哪些？分析时应如何处理？')
ans(doc)
ap(doc, 1, '生物学特征：繁殖能力、生长速率、自然死亡率等内在参数，决定种群自我更新和增长的潜力。')
ap(doc, 2, '环境因素：水温、盐度、溶解氧、营养盐、水流等理化环境条件，直接影响鱼类的生理活动和分布。')
ap(doc, 3, '生物因素：饵料生物丰度、种间竞争、敌害生物等生态相互作用。')
ap(doc, 4, '人为因素：捕捞强度、水利工程建设、水体污染、航运等人为活动的影响日益显著。')
exp(doc, '四种因素相互交织、共同作用，分析时应综合分析、找出主导因素。例如在过度捕捞情形下，即使环境条件适宜，持续的高强度捕捞仍会导致资源衰退。')

doc.save(DOCX)
total=sum(len(p.text) for p in doc.paragraphs)
bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
print(f'文档已保存: {DOCX}')
print(f'总字符数: {total}')
print(f'粗体运行: {bold_c}')
print(f'表格数: {len(doc.tables)}')
