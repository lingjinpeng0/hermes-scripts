#!/usr/bin/env python3
"""
渔业资源学课件整理 — 逐章生成 (Chapter 1 only)
Fixed: Chinese quotes, table captions centered.
"""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = os.path.expanduser(r"~\Downloads\学习\渔业资源学\渔业资源学课件整理.docx")
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

LQ = '\u201c'  # "
RQ = '\u201d'  # "

def sf(r, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False, italic=False):
    r.font.size=sz; r.font.name=en
    r.element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.bold=bold; r.font.italic=italic

def body(doc, txt, indent=True):
    p=doc.add_paragraph(); r=p.add_run(txt); sf(r)
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.first_line_indent=Pt(21) if indent else None

def h1(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 1']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(16),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

def h2(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 2']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(14),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def h3(doc, txt):
    p=doc.add_paragraph(); p.style=doc.styles['Heading 3']
    r=p.add_run(txt); sf(r,cn='黑体',sz=Pt(12),bold=False,italic=False)
    p.paragraph_format.first_line_indent=None; p.paragraph_format.left_indent=None
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def tbl(doc, hd, rows, cap=None):
    if cap:
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pPr=p._p.get_or_add_pPr()
        pJC=OxmlElement('w:jc')
        pJC.set(qn('w:val'),'center')
        pPr.append(pJC)
        r=p.add_run(cap); sf(r)
        p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.first_line_indent=None
    t=doc.add_table(rows=1+len(rows), cols=len(hd))
    t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tl=t._tbl; tp=tl.tblPr
    if tp is None: tp=OxmlElement('w:tblPr'); tl.insert(0,tp)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:w'),'5000'); tw.set(qn('w:type'),'pct'); tp.append(tw)
    tb=OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'nil'); tb.append(b)
    tp.append(tb)
    for ci,h in enumerate(hd):
        c=t.rows[0].cells[ci]; c.text=''; p=c.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Pt(0)
        r=p.add_run(h); sf(r,sz=Pt(10))
        _bdr(c,'12','6')
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if len(str(v))<30 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent=Pt(0)
            r=p.add_run(str(v)); sf(r,sz=Pt(9))
    for ci in range(len(hd)):
        c=t.rows[-1].cells[ci]; tcp=c._tc.get_or_add_tcPr()
        tcb=OxmlElement('w:tcBorders'); bt=OxmlElement('w:bottom')
        bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'12')
        bt.set(qn('w:space'),'0'); bt.set(qn('w:color'),'000000')
        tcb.append(bt); tcp.append(tcb)
    for row in tl.findall(qn('w:tr')):
        trp=row.find(qn('w:trPr'))
        if trp is None: trp=OxmlElement('w:trPr'); row.insert(0,trp)
        old=trp.find(qn('w:cantSplit'))
        if old is not None: trp.remove(old)
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trp.append(cs)
    # make sure table is centered
    tPr=tl.tblPr
    tjc=OxmlElement('w:jc')
    tjc.set(qn('w:val'),'center')
    tPr.append(tjc)

def _bdr(c,ts='12',bs='6'):
    tcp=c._tc.get_or_add_tcPr(); tcb=OxmlElement('w:tcBorders')
    for s,sz in [('top',ts),('bottom',bs)]:
        e=OxmlElement(f'w:{s}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); tcb.append(e)
    tcp.append(tcb)

def q(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'{n}. {txt}'); sf(r,sz=Pt(11))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.first_line_indent=None

def ans(doc):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run('答案:'); sf(r)

def ap(doc, n, txt):
    p=doc.add_paragraph(); r=p.add_run(f'({n}) {txt}'); sf(r,sz=Pt(10))
    p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); ind=OxmlElement('w:ind')
    ind.set(qn('w:left'),'840'); ind.set(qn('w:hanging'),'210'); pPr.append(ind)

def exp(doc, txt):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Pt(21); r=p.add_run(f'解析: {txt}'); sf(r,sz=Pt(10))

# ══════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════
def build_ch1(doc):
    h1(doc, '1. 绪论')

    h2(doc, '1.1 资源的基本概念与分类体系')
    h3(doc, '1.1.1 资源的一般概念')
    body(doc, f'{LQ}资源{RQ}一词在日常生活中广泛应用，其基本含义是指自然界及人类社会中一切能为人类形成资财的要素。《辞海》将其定义为{LQ}资财的来源，一般指自然的财源{RQ}。从广义上看，凡是能够满足人类生产生活需要、具有使用价值或潜在价值的物质、能量、信息乃至人力，均可纳入资源的范畴。根据来源和性质的差异，资源首先可被划分为社会资源与自然资源两大基本类型。社会资源包括人力资源、劳动资源以及经济资源（即人类已经加工过的物质和能量），它们是人类社会活动过程中创造和积累的财富基础。自然资源则是未经人类加工而天然存在的自然物。')
    body(doc, f'自然资源在《辞海》中被定义为{LQ}天然承载的自然物，如土地资源、水资源、生物资源和海洋资源等，不包括人类加工制造的原料，是生产的原料来源和布局场所{RQ}。联合国环境规划署（UNEP）给出了更为系统的定义：在一定时间条件下，能够产生经济价值以提高人类当前和未来福利的自然环境因素的总称。这一定义强调了自然资源的时间性（可利用性随时间和技术条件变化）、经济性（必须能够产生经济价值）和福利性（服务于人类福祉）。')

    h3(doc, '1.1.2 自然资源的分类体系')
    body(doc, '自然资源种类繁多，根据不同的分类标准可划分为多个类别，每一分类维度都揭示着资源的不同属性特征。')
    body(doc, '按存在形态分类：可分为土地资源、气候资源、水资源、矿产资源、生物资源、环境资源（如景观资源、大气资源）等。这种分类直观反映了资源的存在形式，便于从部门管理的角度进行规划利用。')
    body(doc, '按可更新特征分类：这是资源科学中最核心的分类方式之一。可再生资源（如生物资源、水资源、气候资源）具有通过自然循环或生物繁殖而自我更新的能力；非再生资源（如矿产资源、化石燃料）则在地质时间尺度上无法恢复。渔业资源作为生物资源的一种，具有典型的可再生性，但这种可再生性是有条件的——当捕捞强度超过资源的自然补充能力时，资源就会衰退甚至枯竭。')
    body(doc, f'按控制方式分类：可分为专有资源（可由特定主体占有和管理的资源，如养殖水域）和共享资源（如公海渔业资源、大气资源）。共享资源容易引发{LQ}公地悲剧{RQ}（Tragedy of the Commons），即每个个体都追求自身利益最大化，最终导致资源被过度利用乃至毁灭。渔业资源中大量存在共享资源属性，这也是渔业管理面临的根本困境之一。')
    body(doc, '按可持续性分类：包括可持续利用资源（如太阳能、风能）、非持续利用资源（如石油、煤炭）以及介于二者之间的双重性资源。渔业资源属于典型的双重性资源——合理利用可以永续，过度利用则不可持续。按流动性分类：可分为储存性资源（可在原地积累和储存的资源）和流动性资源（如渔业生物资源中的洄游性鱼类，其分布随时间和环境变化而不断移动）。流动性资源的评估和管理比储存性资源复杂得多，因为它涉及空间动态和跨境管理问题。')

    h3(doc, '1.1.3 生物资源与渔业资源')
    body(doc, '生物资源（Biological Resource）是自然资源中最为活跃的一类，包括植物资源、动物资源和微生物资源。按生态系统的不同，又可细分为森林资源、牧草资源、鱼类资源、野生动物资源等。生物资源的本质特征在于其可再生性和可更新性——通过生长、繁殖和代谢，生物资源可以在一定限度内自我更新和增殖。')
    body(doc, '渔业资源（Fishery Resource），亦称水产资源，是生物资源的重要组成部分。关于渔业资源的定义，不同工具书给出了不同的表述：狭义定义指天然水域中具有渔业开发利用价值的鱼类种类和数量；《辞海》定义为水域中蕴藏的各种经济动植物（鱼类、贝类、甲壳类、藻类）的数量；《农业大辞典》与《中国农业百科全书》（水产卷）定义为天然水域中具有开发利用价值的经济动植物种类和数量的总称。按水域类型，渔业资源可分为海洋渔业资源、淡水渔业资源和河口渔业资源三大类。')

    h2(doc, '1.2 渔业资源增殖')
    h3(doc, '1.2.1 增殖的定义与内涵')
    body(doc, '渔业资源增殖（Fishery Resource Enhancement）是指通过人为干预手段增加天然水域中渔业资源种群数量、改善资源结构和水域生态环境的一系列措施。狭义的资源增殖指向天然水域中投放鱼、虾、贝、藻等水生生物幼体（或成体、卵等），以直接增加种群数量，改善和优化水域的渔业资源群落结构。这是最常见、最直接的增殖方式。')
    body(doc, '广义的资源增殖除人工放流外，还包括改善水域生态环境（如水质修复、产卵场保护）、向特定水域投放人工装置（如附卵器、人工鱼礁）、野生种群的保护和病害防治等间接增加水域种群资源量的措施。广义的增殖概念已与渔业管理的理念高度融合，强调生态系统的整体修复和可持续性。')

    h3(doc, '1.2.2 增殖的主要方法')
    body(doc, '（1）人工放流（Stock Enhancement）是渔业资源增殖最核心的手段，即在科学评估的基础上，有计划地向天然水域投放人工培育的苗种，使其在自然环境中生长、发育，最终补充或恢复目标种群的资源量。中国是世界上开展人工放流规模最大的国家之一，每年在黄海、东海、南海及各大淡水流域放流大量经济鱼类、虾蟹类和贝类苗种。')
    body(doc, f'放流与放生有本质区别：放流是建立在科学调查和评估基础之上的系统性资源管理行为——放流前需要评估水域生态容量、目标种群的补充状况、遗传多样性保护以及放流后的效果跟踪。而放生多为个人或团体的自发行为，往往缺乏科学依据。科学放流与盲目放生有着本质区别，前者是资源管理的重要手段，后者则可能带来生态风险。')
    body(doc, '（2）引种驯化（Introduction and Acclimation）是指将某一水域的经济物种引入到新的水域环境中，使其适应新环境并建立可自我维持的自然种群。成功的引种驯化可以丰富目标水域的生物多样性和经济产出。然而，引种驯化也存在生态风险，引入种可能与本地物种竞争资源、捕食本地种或杂交导致遗传污染，因此引种前必须进行严格的生态风险评估。')
    body(doc, '（3）改善水域环境是资源增殖的间接但重要的手段。包括：修复鱼类产卵场和栖息地；控制工业和生活污水排放，改善水质；恢复水生植被，为鱼类提供栖息和索饵场所；拆除或改造阻隔鱼类洄游的水利设施，恢复河湖连通性。')

    h3(doc, '1.2.3 人工鱼礁')
    body(doc, '人工鱼礁是渔业资源增殖中极具特色的一项措施，其灵感源于1860年美国一场洪水将大树冲进海湾后，枝干上附着生长了诸多海洋生物，大量鱼类随之而来的自然现象。人工鱼礁按水层位置可分为沉式鱼礁和浮式鱼礁；按功能可分为养殖型鱼礁、幼鱼保护型鱼礁、增殖型鱼礁、渔获型鱼礁和游钓型鱼礁。人工鱼礁的原理是通过投放人工构造物改变海底地形和水流条件，为鱼类提供栖息、避难和索饵场所，从而聚集和增殖渔业资源。')

    tbl(doc,
        ['鱼礁类型', '功能特点', '应用场景'],
        [
            ['养殖型鱼礁', '为养殖对象提供附着基和栖息环境', '海参、鲍鱼、龙虾等底栖经济种养殖'],
            ['幼鱼保护型鱼礁', '提供隐蔽场所，降低被捕食风险', '近岸幼鱼栖息地保护'],
            ['增殖型鱼礁', '创造新栖息地，增加资源量', '资源衰退海域的增殖修复'],
            ['渔获型鱼礁', '聚集鱼类，提高捕捞效率', '配合选择性捕捞作业'],
            ['游钓型鱼礁', '聚集休闲性鱼种', '休闲渔业和生态旅游'],
        ],
        cap='表1-1.人工鱼礁的主要类型')

    h2(doc, '1.3 水产养殖与捕捞渔业')
    body(doc, '水产养殖（Aquaculture）是指在人为控制下繁殖、培育和收获水生动植物的生产活动，一般包括在人工饲养管理下从苗种养成水产品的全过程，广义上也可包括水产资源增殖。主要养殖方式包括网箱养殖、工厂化养殖和池塘养殖。中国是世界第一水产养殖大国，养殖产量占全球养殖产量的60%以上。')

    tbl(doc,
        ['指标', '2024年产量(万吨)', '同比增长(%)', '占比(%)'],
        [
            ['水产品总产量', '7357.59', '3.39', '100.0'],
            ['养殖产量', '6060.03', '4.31', '82.4'],
            ['捕捞产量', '1297.56', '-0.69', '17.6'],
            ['海水产品', '3708.87', '3.45', '50.4'],
            ['淡水产品', '3648.72', '3.34', '49.6'],
        ],
        cap='表1-2.2024年全国水产品产量概况')

    body(doc, '2024年全国水产品总产量7357.59万吨，其中养殖产量6060.03万吨（占82.4%），捕捞产量1297.56万吨（占17.6%），养殖产品与捕捞产品的产量比例为82.4:17.6。远洋渔业产量218.91万吨，同比下降5.74%。全国水产品人均占有量52.25千克，比上年增加1.77千克。海水产品产量3708.87万吨，淡水产品产量3648.72万吨，海水产品与淡水产品的产量比例约为50.4:49.6。从产业结构看，养殖产量占比超过80%，表明中国已成为以养殖为主的水产大国。')

    h2(doc, '1.4 渔业资源学的定义与研究内容')
    h3(doc, '1.4.1 学科定义')
    body(doc, '渔业资源学属于渔业科学和生态学相关应用科学的范畴，研究的中心内容是渔业生物资源群体的变动规律，主要目标是为渔业水域及生物资源的持续利用、渔业生产的健康发展提供科学依据。使用过的名称包括：水产资源学、水产资源力学、水产资源生物学、渔业科学、渔业生物学、资源生物学等。')

    h3(doc, '1.4.2 主要研究内容')
    body(doc, '渔业资源学的研究内容主要包括六个方面：(1)摸清渔业生物的种群结构、早期生活史和繁殖、摄食、生长、死亡、补充、洄游、分布等生物学特性；(2)种群数量动态规律的研究，包括资源量变动的原因、补充机制、资源量的评估方法；(3)渔业水域栖息地生态环境与早期发育阶段的关系及其影响机制，研究水文、气象等环境因素与渔业资源集群、洄游分布的关系；(4)研究捕捞方式、强度对种群数量、年龄组成、补充和群落结构及其动态特征的作用和影响机制；(5)渔业水域生态系统的结构与功能的研究，以保持渔业生物群落结构的多样性和稳定性为基础；(6)渔业资源的开发、管理和增殖涉及到资源分享、就业、投入和产出等经济社会诸多方面的平衡和矛盾，不能单从生物学因素来考虑渔业问题。')

    h2(doc, '1.5 渔业资源学的起源与发展')
    h3(doc, '1.5.1 国际发展历程')
    body(doc, f'现代渔业资源研究始于欧洲北海。19世纪末，德国学者Heincke是大西洋鲱生活史研究的开创者，对鲱生活史和鲽类生长与密度关系进行了深入调查研究。Heincke在1898年首先把渔业生物学和捕捞统计资料结合起来研究种群数量动态，并据此提出了{LQ}繁殖论{RQ}，认为由于种群分布区域的局限性，种群数量和年龄组成的变化主要取决于捕捞强度，捕捞强度加大必然导致渔获量下降、鱼体小型化和低龄化，提出了限定最小可捕长度的方法保护资源。')
    body(doc, '代表性人物及著作包括：Russel（1932）《渔业研究》、Dymond（1948）《鱼类种群研究》、相川秋（1941）《水产资源学》、Gulland（1977）《鱼类种群数量变动》和《鱼类资源评估》、蒙纳斯蒂尔斯基（1952）《鱼类种群数量变动》等。1902年成立于丹麦哥本哈根的ICES（国际海洋考察理事会）是协调促进海洋科学考察的重要国际组织，其任务包括促进和鼓励为研究北大西洋和邻近海域的海洋环境及其生物资源进行考察和调查，并向各国政府、区域渔业管理部门及污染防治委员会提供科学信息和建议。')

    h3(doc, '1.5.2 中国研究概况')
    body(doc, '中国渔业资源研究历史悠久。早在3000年前就有关于鱼的文字记载；周代战国时期渔具渔法多样化，为保护资源规定了禁渔期和网目尺寸，并形成了渔业管理队伍；明代沿海渔民对大、小黄鱼的生活习性和洄游路线有了深入了解。20世纪50年代末60年代初，各省市科研机构对河流、湖泊进行渔业资源和生态环境调查。80年代初全国同步开展了内陆水域四大水系及重要湖泊、水库的渔业资源调查和区划研究，查清了青鱼、草鱼、鲢、鳙、鲤、鲫、鳊、鲌、鳡、银鱼、江鳕、乌鳢、鳜、黄鳝、中华绒螯蟹等主要经济鱼类和水生生物的生物学特性。')
    body(doc, '70年代中期之前，以渔场调查和捕捞对象的渔业生物学、洄游分布、数量变动和资源评估为中心的种群动态研究为主。70年代中期之后，则以近海渔业资源保护、管理和增殖以及持续利用为中心，侧重于渔场生态环境、初级生产力、饵料基础、种间关系、渔业资源群落结构、补充特性和机制以及年间变化、资源增殖潜力和生态容量研究。')

    h2(doc, '1.6 发展趋势与面临的问题')
    h3(doc, '1.6.1 淡水渔业资源面临的主要问题')
    body(doc, f'淡水渔业资源面临的主要问题包括：(1)鱼类群落组成变化，经济鱼种比例下降，低值鱼种比例上升；(2)资源小型化，低龄鱼比例增加，大型个体减少；(3)水域环境质量下降，富营养化和污染问题突出；(4)河湖连通性受阻，水利工程建设阻断了鱼类洄游通道；(5)水资源量减少，湖泊萎缩、河流断流；(6)捕捞网具滥用，{LQ}绝户网{RQ}等非法渔具严重破坏资源。')

    h3(doc, '1.6.2 发展趋势')
    body(doc, '渔业资源学的发展趋势包括：(1)从单鱼种研究扩展到生态系统层面的多物种管理——即基于生态系统的渔业管理（EBFM）；(2)先进技术的深度融合——声学资源评估技术、卫星遥感、人工智能和大数据分析在渔业资源调查和评估中的广泛应用；(3)绿色可持续发展模式的构建——发展负责任捕捞和生态养殖，实现渔业资源的可持续利用；(4)国际合作在渔业管理中的重要性日益凸显——跨界鱼类种群的管理需要国际社会的共同努力。')

    h2(doc, '1.7 思考题')

    q(doc, 1, '长江十年禁捕对水产行业及渔业资源保护和利用有何影响？')
    ans(doc)
    ap(doc, 1, '对水产行业的影响：禁捕倒逼水产养殖业转型升级，推动养殖技术提升和产业结构优化，增加了对优质苗种、配合饲料和养殖装备的需求，促进了工厂化养殖和智慧渔业等现代养殖模式的发展。')
    ap(doc, 2, '对渔业资源保护：禁捕有效降低了捕捞强度，为鱼类等水生生物提供了繁殖和恢复的时间与空间。自2021年启动以来，长江水生生物多样性恢复向好，江豚栖息范围扩大，珍稀鱼类种群数量明显回升。')
    ap(doc, 3, f'对资源利用：倒逼利用方式从{LQ}数量型{RQ}转向{LQ}质量效益型{RQ}，促进渔业资源的科学管理和可持续利用，同时推动渔民转产转业和生态补偿机制的建立。')
    exp(doc, '长江十年禁渔是习近平生态文明思想的重要实践，对长江经济带生态优先、绿色发展具有深远意义。')

    q(doc, 2, '渔业资源对人类生存和发展有何影响？')
    ans(doc)
    ap(doc, 1, '提供优质动物蛋白：渔业资源是人类优质蛋白质的重要来源，对保障粮食安全和国民营养健康具有不可替代的作用。2024年全国水产品人均占有量52.25千克。')
    ap(doc, 2, '促进经济发展：渔业及上下游产业为沿海和内陆地区提供了大量就业机会和经济效益。中国是世界第一渔业大国，水产品出口创汇在农业贸易中占有重要地位。')
    ap(doc, 3, '维持生态平衡：渔业资源是水生生态系统的重要组成部分，对维持水域食物网稳定和生态平衡具有关键作用。')
    ap(doc, 4, '文化价值：渔业活动具有丰富的文化内涵，涉及传统文化传承、休闲渔业、生态旅游等多个方面。')
    exp(doc, '渔业资源的可持续利用关乎人类社会的可持续发展，需要平衡开发利用与保护管理的关系。')

    q(doc, 3, '简述渔业资源增殖的狭义和广义定义及其主要方法。')
    ans(doc)
    ap(doc, 1, '狭义定义：向天然水域中投放鱼、虾、贝、藻等水生生物幼体（或成体、卵等）以增加种群数量，改善和优化水域的渔业资源群落结构。')
    ap(doc, 2, '广义定义：除人工放流外，还包括改善水域的生态环境（水质修复、产卵场保护）、向水域投放人工装置（附卵器、人工鱼礁）以及野生种群的防治保护等间接措施。')
    ap(doc, 3, f'主要方法：人工放流（核心手段）、引种驯化、改善水域环境、建设人工鱼礁。需注意科学{LQ}放流{RQ}与盲目{LQ}放生{RQ}的本质区别。')
    exp(doc, '科学放流应基于种群动力学和生态承载力评估，做到针对性增殖，避免盲目放流造成生态风险。')

# ══════════════════════════════════
# MAIN
# ══════════════════════════════════
def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name='Times New Roman'; style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    style.font.size=Pt(10.5); style.font.bold=False; style.font.italic=False
    style.paragraph_format.line_spacing=1.2; style.paragraph_format.first_line_indent=Pt(21)
    for lv in [1,2,3]:
        s=doc.styles[f'Heading {lv}']
        s.font.name='Times New Roman'; s.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
        s.font.bold=False; s.font.italic=False; s.font.color.rgb=None
    build_ch1(doc)
    doc.save(OUTPUT)
    total=sum(len(p.text) for p in doc.paragraphs)
    bold_c=sum(1 for p in doc.paragraphs for r in p.runs if r.font.bold)
    print(f'文档已生成: {OUTPUT}')
    print(f'总字符数: {total}')
    print(f'段落数: {len(doc.paragraphs)}')
    print(f'表格数: {len(doc.tables)}')
    print(f'粗体运行: {bold_c}')
    # Check for escaped unicode in output
    for i,p in enumerate(doc.paragraphs):
        if '\\u201c' in p.text or '\\u201d' in p.text:
            print(f'WARNING: P{i} has literal \\u201c/h escape: {p.text[:60]}')

if __name__=='__main__':
    main()
