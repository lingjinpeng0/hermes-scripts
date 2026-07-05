"""Verify ch0+ch1 document."""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'C:\Users\Rei\Downloads\学习\生物统计\生物统计课件整理.docx')

print("=== 段落概览 ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if not text:
        continue

    has_bold = any(r.font.bold for r in p.runs if r.font.bold is not None)
    has_italic = any(r.font.italic for r in p.runs if r.font.italic is not None)
    fw = sum(1 for c in text if 0xFF00 <= ord(c) <= 0xFFEF)
    fi = p.paragraph_format.first_line_indent

    flag = ''
    if has_bold: flag += ' [B]'
    if has_italic: flag += ' [I]'
    if fw > 0: flag += f' [FW:{fw}]'

    short = text[:90]
    print(f'[{i:3d}] {style:15s} fi={str(fi):8s}{flag} | {short}')

print()
print("=== 汇总 ===")
print(f'段落: {len(doc.paragraphs)}')
print(f'表格: {len(doc.tables)}')
all_text = ''.join(p.text for p in doc.paragraphs)
fw = sum(1 for c in all_text if 0xFF00 <= ord(c) <= 0xFFEF)
print(f'全角标点: {fw}')
h2 = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 2')
h3 = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 3')
h4 = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 4')
print(f'Heading 2: {h2}, Heading 3: {h3}, Heading 4: {h4}')

# Check for bold/italic
bold_any = False
italic_any = False
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.bold: bold_any = True
        if r.font.italic: italic_any = True
print(f'有粗体: {bold_any}')
print(f'有斜体: {italic_any}')

# Check tables
for ti, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    print(f'表{ti}: {rows}行 x {cols}列')
