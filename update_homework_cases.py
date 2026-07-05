"""Update 推文作业答案汇总 - insert new domestic cases after 京东物流 case"""
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn

INPUT = r'C:\Users\Rei\Downloads\学习\智慧渔业\汇报\智慧渔业推文作业答案汇总.docx'

doc = Document(INPUT)

def sf(run, cn='宋体', en='Times New Roman', sz=Pt(10.5)):
    run.font.name = en; run.font.size = sz; run.font.bold = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def make_body_para(text, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text); sf(r)
    p.paragraph_format.first_line_indent = Pt(21) if indent else None
    p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(3)
    return p

def make_bullet(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f'({num}) {text}'); sf(r)
    p.paragraph_format.first_line_indent = None; p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(2)
    return p

def make_sub(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f'({num}) '); sf(r)
    r2 = p.add_run(text); sf(r2, sz=Pt(10.5))
    p.paragraph_format.first_line_indent = None; p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(2)
    return p

# Find the paragraph after 京东 (paragraph 20) - we need to insert after it
# In python-docx, we can't easily insert after a specific paragraph.
# Instead, we insert at the end, then move elements in XML.

# Find the anchor paragraph (京东's paragraph)
anchor_text = '京东物流建立的"三位一体"综合冷链基础网络体系'
insert_after_idx = None
for i, p in enumerate(doc.paragraphs):
    if anchor_text in p.text:
        insert_after_idx = i
        break

if insert_after_idx is None:
    print('ERROR: Could not find anchor paragraph')
    exit(1)

print(f'Found anchor at paragraph {insert_after_idx}')

# Create new paragraphs
new_paras = []

# (4) 安徽明光
bp = make_body_para('')
new_paras.append(('bullet', '(4) 安徽明光数字孪生智慧渔场'))
new_paras.append(('body', '安徽省明光市依托国家智慧渔业创新应用项目(总投资5339万元), 融合GIS、数字孪生、智能装备、物联网和AI大数据技术, 构建覆盖主要淡水养殖场景的数字化融合新模式。项目部署软件系统20套, 联动控制硬件设施超1000台(套), 建成明光市数字渔业大数据指挥调度中心, 集成19个子系统覆盖生产、管理与服务三大板块。创新采用"水陆空"三维协作模式——无人机空中巡查、无人船水面投喂、水下ROV巡检监测, 结合图像识别算法自动识别网衣破损、附着物、鱼类病害及逃逸迹象, 实现"数字孪生"渔场管理与"智能协同"精准作业。'))

# (5) 连江闽投1号
new_paras.append(('bullet', '(5) 福建连江"闽投1号"深远海养殖平台'))
new_paras.append(('body', '全国首个半潜式渔旅融合平台"闽投1号"部署在福建连江定海湾, 配备智慧渔业系统, 使大黄鱼拥有"智能家居"。智慧海洋渔业管理平台实时收集水质、水温、盐度、海流等海洋环境数据, 通过鱼脸识别、水下动感捕捉摄像系统实现精准投喂和智能清洁。搭载水下巡检机器人深入水底巡查养殖网衣, 网箱清洁机器人自动清理藤壶等附着物。福建首个"定制化海上生鲜无人机安全速达"应用场景在该平台启用, 大黄鱼可在50分钟内从深海"飞"到福州市区。连江已投放11台桁架类深远海养殖平台, 数量位居全国县级第一, 建成全国规模最大的海上智慧牧场。'))

# (6) 渔军师平台
new_paras.append(('bullet', '(6) 机智云"渔军师"物联网智慧渔业大数据平台'))
new_paras.append(('body', '广州机智云依托自身工业互联网平台, 打造"渔军师"智慧渔业解决方案, 通过智能增氧机、鱼料控制器、荧光法溶氧仪、pH检测仪、AI摄像头等设备, 对养殖池塘环境、水质变化、鱼群状态等核心数据进行实时采集与上云。平台融合大数据与AI技术, 对水质、温度、溶解氧等关键指标进行动态分析, 预测水体变化趋势、鱼类生长周期和病害风险, 输出智能化调控建议, 辅助制定精准的投喂计划和防病机制。同时推动养殖户、科研院所、政府监管等多方数据共享与业务协同, 已在多个淡水养殖示范区落地应用。'))

# Now we need to insert these paragraphs AFTER the anchor paragraph
# Get the body element
body = doc.element.body

# Find the anchor element
anchor_elem = None
for p_elem in body:
    if p_elem.tag == qn('w:p'):
        # Check if this paragraph contains anchor text
        texts = [t.text for t in p_elem.iter(qn('w:t')) if t.text]
        full_text = ''.join(texts)
        if anchor_text in full_text:
            anchor_elem = p_elem
            break

if anchor_elem is None:
    print('ERROR: Could not find anchor element in XML')
    exit(1)

print(f'Found anchor element, inserting {len(new_paras)} new elements')

# Insert after the anchor
insert_pos = list(body).index(anchor_elem)
insert_idx = insert_pos + 1

for ptype, text in new_paras:
    if ptype == 'bullet':
        p = make_bullet('', text)  # We'll fix the number
    else:
        p = make_body_para(text)
    
    # Fix bullet numbering - the text already has the number
    if ptype == 'bullet':
        p.paragraphs[0].clear()
        r = p.paragraphs[0].add_run(text); sf(r)
    
    # Move the paragraph element to the right position
    p_elem = p._p
    body.remove(p_elem)
    
    # Insert at the right position
    insert_idx += 1
    body.insert(insert_idx, p_elem)

doc.save(INPUT)
print('Done! Document updated.')
