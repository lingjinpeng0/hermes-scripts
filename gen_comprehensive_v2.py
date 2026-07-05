"""智慧渔业综合备考文档 v2 - 直接从源文件拼接内容"""
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = r'C:\Users\Rei\Downloads\学习\智慧渔业\智慧渔业综合备考文档.docx'

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10.5); style.font.bold = False
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format; pf.line_spacing = 1.2; pf.first_line_indent = Pt(21); pf.space_after = Pt(3)

def sf(run, cn='宋体', en='Times New Roman', sz=Pt(10.5)):
    run.font.name = en; run.font.size = sz; run.font.bold = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def add_heading(text, level=2):
    p = doc.add_paragraph()
    if level == 0:
        r = p.add_run(text); sf(r, sz=Pt(16))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None; p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p.style = doc.styles['Heading 2']; r = p.add_run(text); sf(r, cn='黑体', sz=Pt(14))
        p.paragraph_format.first_line_indent = None; p.paragraph_format.left_indent = None; p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.style = doc.styles['Heading 3']; r = p.add_run(text); sf(r, cn='黑体', sz=Pt(12))
        p.paragraph_format.first_line_indent = None; p.paragraph_format.left_indent = None; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    return p

def add_body(text):
    if not text.strip():
        return
    p = doc.add_paragraph(); r = p.add_run(text); sf(r)
    p.paragraph_format.first_line_indent = Pt(21); p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(3)

def add_bni(text):
    if not text.strip():
        return
    p = doc.add_paragraph(); r = p.add_run(text); sf(r)
    p.paragraph_format.first_line_indent = None; p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(2)

def add_qa_q(text):
    p = doc.add_paragraph(); r = p.add_run(text); sf(r, sz=Pt(11))
    p.paragraph_format.first_line_indent = None; p.paragraph_format.space_before = Pt(8); p.paragraph_format.line_spacing = 1.2

def add_qa_a():
    p = doc.add_paragraph(); r = p.add_run('答案:'); sf(r)
    p.paragraph_format.first_line_indent = Pt(21); p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(3)

def add_qa_p(num, text):
    p = doc.add_paragraph(); r = p.add_run(f'({num}) {text}'); sf(r)
    p.paragraph_format.first_line_indent = None; p.paragraph_format.left_indent = Twips(840); p.paragraph_format.hanging_indent = Twips(210)
    p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(2)

def add_qa_parse(text):
    p = doc.add_paragraph(); r = p.add_run(f'解析: {text}'); sf(r)
    p.paragraph_format.first_line_indent = Pt(21); p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(8)

# Read source texts
base = os.path.expanduser('~/.hermes/project-data/智慧渔业/source_texts')
sources = {}
for key in ['notes', 'focus', 'qa', 'hw']:
    with open(os.path.join(base, f'{key}.txt'), 'r', encoding='utf-8') as f:
        sources[key] = f.read().split('\n')

# Parse into chapter sections
def get_chapter_lines(source_list, ch_markers):
    """Get lines belonging to a chapter"""
    result = []
    in_ch = False
    for line in source_list:
        if not line.strip():
            continue
        if any(marker in line for marker in ch_markers):
            in_ch = True
            result.append(line)
        elif in_ch:
            # Check if this is a new chapter heading (H2)
            if line.startswith('[Heading 2]') and any(str(n) in line for n in range(1,13)):
                # Check if it's a chapter we want
                is_new_ch = True
                for marker in ch_markers:
                    if marker in line:
                        is_new_ch = False
                        break
                if is_new_ch:
                    break
            result.append(line)
    return result

def extract_body_lines(chapter_lines):
    """Extract body text paragraphs from chapter lines"""
    bodies = []
    for line in chapter_lines:
        if line.startswith('[Normal]') or line.startswith('[Heading'):
            # Remove prefix
            text = line.split('] ', 1)[1] if '] ' in line else line
            bodies.append((line.split(']')[0].lstrip('['), text))
    return bodies

# Title
add_heading('智慧渔业综合备考文档', 0)
add_body('课程: 智慧渔业 | 华中农业大学水产学院')
add_body('考试范围: 第2-6章, 第9-10章(第1章不考)')
add_body('内容说明: 每章整合课件笔记、考试重点和Q&A练习。末尾附综合应用题和推文作业答案汇总。')

# ===== Chapter sections =====
# Map chapters to their markers in source texts
chapters = [
    ('2-3章 感觉器官与行为(合并)', ['2.', '第2章', '第2-3章', '2. 鱼类']),
    ('第4章 鱼类群体行为学', ['4.', '第4章', '4. 鱼类群体']),
    ('第5章 渔业物联网技术', ['5.', '第5章', '5. 渔业物联网']),
    ('第6章 渔业大数据', ['6.', '第6章', '6. 渔业大数据']),
    ('第9章 智能渔业装备', ['9.', '第9章', '9. 智能渔业']),
    ('第10章 智能陆基循环水养殖系统(RAS)', ['10.', '第10章', '10. 智能陆基']),
]

q_num = 0

for ch_title, markers in chapters:
    add_heading(ch_title, 1)
    
    # 1. Notes content
    ch_notes = get_chapter_lines(sources['notes'], markers)
    if ch_notes:
        for style, text in extract_body_lines(ch_notes):
            if style in ('Heading 2', 'Heading 3'):
                # Skip the chapter title itself (already added) and 思考题
                if any(kw in text for kw in ['思考题', '答案:', '解析:', '1. 绪论']):
                    continue
                add_heading(text, 2 if style == 'Heading 2' else 2)
            elif text.startswith('答案:'):
                # Skip answers in notes (they're basic)
                pass
            elif text.startswith('(') and text[1:2].isdigit():
                add_bni(text)
            else:
                add_body(text)
    
    # 2. Focus content (only key points)
    ch_focus = get_chapter_lines(sources['focus'], markers)
    if ch_focus:
        add_heading('考试要点', 2)
        for style, text in extract_body_lines(ch_focus):
            if '重点掌握' in text or '重点掌握' in text:
                add_body(text)
            elif text.startswith('(') and text[1:2].isdigit():
                add_bni(text)
    
    # 3. Q&A content
    ch_qa = get_chapter_lines(sources['qa'], markers)
    if ch_qa:
        add_heading('Q&A练习', 2)
        for style, text in extract_body_lines(ch_qa):
            if text.startswith('Q') and '. ' in text[:5]:
                q_num += 1
                add_qa_q(text)
            elif text == '答案:':
                add_qa_a()
            elif text.startswith('(') and text[1:2].isdigit():
                add_qa_p(text[1], text[4:])  # Extract number and content
            elif text.startswith('解析:'):
                add_qa_parse(text[3:])

# ===== 应用题 =====
ch_qa_app = get_chapter_lines(sources['qa'], ['应用题'])
if ch_qa_app:
    add_heading('应用题', 1)
    for style, text in extract_body_lines(ch_qa_app):
        if text.startswith('Q') and '. ' in text[:5]:
            add_qa_q(text)
        elif text == '答案:':
            add_qa_a()
        elif text.startswith('(') and text[1:2].isdigit():
            add_qa_p(text[1], text[4:])

# ===== 附录: 推文作业答案 =====
add_heading('附录: 推文作业答案汇总', 1)

hw_full = sources['hw']
hw_started = False
for line in hw_full:
    style = line.split(']')[0].lstrip('[')
    text = line.split('] ', 1)[1] if '] ' in line else line
    
    if text.startswith('智慧渔业推文作业答案汇总') or text.startswith('课程:'):
        continue
    
    if style == 'Heading 2' or style == 'Heading 3':
        add_heading(text, 1 if style == 'Heading 2' else 2)
    elif text.startswith('解析:'):
        add_qa_parse(text[3:])
    elif text.startswith('答案:'):
        pass  # Skip answer labels in appendix
    elif text.startswith('(') and text[1:2].isdigit():
        add_bni(text)
    else:
        add_body(text)

doc.save(OUTPUT)
print(f'已生成: {OUTPUT}')
