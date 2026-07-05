"""
Generate the complete 智慧渔业 courseware summary docx from organized text content.
Reads 3 content files, combines them, and generates a single well-formatted Word document.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_styles(doc):
    """Set up document styles according to formatting spec."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.first_line_indent = Pt(21)
    style.paragraph_format.line_spacing = 1.2
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Heading 1 - Main title (三号 16pt)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.first_line_indent = None
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2 - Chapter title (四号 14pt)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = False
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = None
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3 - Section title (小四 12pt)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = None
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)

    # List Bullet style
    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Times New Roman'
        lb.font.size = Pt(10.5)
        lb.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        lb.paragraph_format.first_line_indent = Pt(0)
        lb.paragraph_format.line_spacing = 1.2
    except:
        pass


def add_paragraph(doc, text, style='Normal', bold=False):
    """Add a paragraph with proper formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    if bold:
        run.font.bold = False  # No bold per spec
    return p


def add_heading(doc, text, level=2):
    """Add a heading with proper formatting."""
    p = doc.add_heading(text, level=level)
    # Remove bold from heading runs
    for run in p.runs:
        run.font.bold = False
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_three_line_table(doc, headers, rows, caption):
    """Add a three-line table with caption above."""
    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.first_line_indent = Pt(0)
    run = cap_p.add_run(caption)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = header
        for p in cell.paragraphs:
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(9)

    # Apply three-line table borders
    apply_three_line_borders(table)

    # Set table width to full page
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    old_w = tbl_pr.find(qn('w:tblW'))
    if old_w is not None:
        tbl_pr.remove(old_w)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '5000')
    tbl_w.set(qn('w:type'), 'pct')
    tbl_pr.append(tbl_w)

    # Prevent rows from splitting across pages
    for row in tbl.findall(qn('w:tr')):
        tr_pr = row.find(qn('w:trPr'))
        if tr_pr is None:
            tr_pr = OxmlElement('w:trPr')
            row.insert(0, tr_pr)
        cant = OxmlElement('w:cantSplit')
        cant.set(qn('w:val'), '1')
        tr_pr.append(cant)

    # Add spacing after table
    doc.add_paragraph()

    return table


def apply_three_line_borders(table):
    """Apply three-line table borders."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Remove table-level borders
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # Apply cell-level borders
    rows = tbl.findall(qn('w:tr'))
    for i, row in enumerate(rows):
        cells = row.findall(qn('w:tc'))
        for cell in cells:
            tc_pr = cell.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                cell.insert(0, tc_pr)

            tc_borders = OxmlElement('w:tcBorders')

            # Top border
            top = OxmlElement('w:top')
            if i == 0:  # Header row top - thick
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')
            else:
                top.set(qn('w:val'), 'nil')
            tc_borders.append(top)

            # Bottom border
            bottom = OxmlElement('w:bottom')
            if i == 0:  # Header row bottom - thin
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
            elif i == len(rows) - 1:  # Last row bottom - thick
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
            else:
                bottom.set(qn('w:val'), 'nil')
            tc_borders.append(bottom)

            # Left and right borders - nil
            for side in ['left', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tc_borders.append(border)

            old_borders = tc_pr.find(qn('w:tcBorders'))
            if old_borders is not None:
                tc_pr.remove(old_borders)
            tc_pr.append(tc_borders)


def parse_content(content):
    """Parse organized text content into structured chapters."""
    chapters = []
    # Split by chapter breaks
    parts = content.split('===CHAPTER_BREAK===')

    for part in parts:
        part = part.strip()
        if not part:
            continue

        lines = part.split('\n')
        # First non-empty line should be the chapter title
        title = ''
        body_lines = []
        found_title = False

        for line in lines:
            if not found_title and line.strip():
                title = line.strip()
                found_title = True
            else:
                body_lines.append(line)

        if title:
            chapters.append({
                'title': title,
                'body': '\n'.join(body_lines)
            })

    return chapters


def process_chapter_body(doc, body):
    """Process chapter body text, identifying headings, paragraphs, tables, and questions."""
    lines = body.split('\n')
    i = 0
    table_num = 1

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Detect section headings (lines like "一、..." or "第一节..." or "1. ..." or "1.1 ...")
        if re.match(r'^[一二三四五六七八九十]+[、.]', line):
            add_heading(doc, line, level=3)
            i += 1
            continue

        if re.match(r'^第[一二三四五六七八九十]+节', line):
            add_heading(doc, line, level=3)
            i += 1
            continue

        # Detect markdown tables
        if line.startswith('|') and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            rows = []
            for tl in table_lines:
                if re.match(r'^\|[\s\-:]+\|', tl):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)

            if rows:
                headers = rows[0]
                data = rows[1:] if len(rows) > 1 else []
                # Generate table caption
                # Look back for a description line
                caption = f"表{table_num}"
                add_three_line_table(doc, headers, data, caption)
                table_num += 1
            continue

        # Detect thinking questions / homework
        if re.match(r'^(思考题|课后习题|课后作业|第二次作业)', line):
            add_heading(doc, line, level=3)
            i += 1
            continue

        # Detect numbered questions (1. xxx)
        if re.match(r'^\d+\.\s', line):
            p = add_paragraph(doc, line, style='Normal')
            p.paragraph_format.first_line_indent = Pt(0)
            i += 1
            continue

        # Regular paragraph - accumulate until next heading or empty line
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            if re.match(r'^[一二三四五六七八九十]+[、.]', next_line):
                break
            if re.match(r'^第[一二三四五六七八九十]+节', next_line):
                break
            if re.match(r'^(思考题|课后习题|课后作业|第二次作业)', next_line):
                break
            if next_line.startswith('|'):
                break
            if re.match(r'^\d+\.\s', next_line) and len(para_lines) > 0:
                break
            para_lines.append(next_line)
            i += 1

        text = ' '.join(para_lines)
        if text:
            add_paragraph(doc, text, style='Normal')


def generate_document(content_files, output_path):
    """Generate the complete document."""
    # Read and combine all content
    all_content = []
    for f in content_files:
        with open(f, 'r', encoding='utf-8') as fh:
            all_content.append(fh.read())

    combined = '\n===CHAPTER_BREAK===\n'.join(all_content)

    # Parse into chapters
    chapters = parse_content(combined)

    print(f"Parsed {len(chapters)} chapters")
    for ch in chapters:
        print(f"  - {ch['title']}: {len(ch['body'])} chars")

    # Create document
    doc = Document()
    setup_styles(doc)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run('智慧渔业')
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run('课件整理')
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('华中农业大学')
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run('参考教材: 《智慧渔业》, 高坚主编, 高等教育出版社, 2025')
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # Process each chapter
    for ch_idx, ch in enumerate(chapters):
        # Chapter heading
        add_heading(doc, ch['title'], level=2)

        # Process body content
        process_chapter_body(doc, ch['body'])

        # Add page break after each chapter (except last)
        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    # Save
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == '__main__':
    content_files = [
        r"C:\Users\Rei\.hermes\temp\ch01_04_content.txt",
        r"C:\Users\Rei\.hermes\temp\ch05_08_content.txt",
        r"C:\Users\Rei\.hermes\temp\ch09_12_content.txt"
    ]
    output = r"C:\Users\Rei\Downloads\学习\智慧渔业\智慧渔业课件整理.docx"
    generate_document(content_files, output)
